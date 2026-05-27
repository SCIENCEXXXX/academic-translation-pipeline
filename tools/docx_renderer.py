from __future__ import annotations

from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from .extract_pdf_images import ExtractedImage
    from .layout_analyzer import RebuiltPage
except ImportError:
    from extract_pdf_images import ExtractedImage
    from layout_analyzer import RebuiltPage


BODY_FONT = "宋体"
HEADING_FONT = "黑体"
ASCII_FONT = "Arial"


def set_run_font(run, east_asian: str = BODY_FONT, ascii_font: str = ASCII_FONT, size: Pt | None = None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asian)
    if size is not None:
        run.font.size = size


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def set_paragraph_indent(paragraph, first_line_chars: float = 2.0) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(10.5 * first_line_chars)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = ASCII_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    style_map = {
        "Heading 1": (HEADING_FONT, Pt(16)),
        "Heading 2": (HEADING_FONT, Pt(12)),
        "Heading 3": (HEADING_FONT, Pt(12)),
    }
    for style_name, (font_name, size) in style_map.items():
        style = styles[style_name]
        style.font.name = ASCII_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = size
        style.font.bold = True


def add_title_page(document: Document, title: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run(title)
    set_run_font(run, HEADING_FONT, ASCII_FONT, Pt(22))
    run.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("中文译制排版稿")
    set_run_font(run, BODY_FONT, ASCII_FONT, Pt(14))

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    run = p.add_run("供校对、排版与出版前审读使用")
    set_run_font(run, BODY_FONT, ASCII_FONT, Pt(10.5))
    run.font.color.rgb = RGBColor(100, 100, 100)
    document.add_section(WD_SECTION.NEW_PAGE)


def add_source_page_marker(document: Document, page_number: int) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"—— 原书第 {page_number} 页 ——")
    set_run_font(run, BODY_FONT, ASCII_FONT, Pt(10.5))
    run.font.color.rgb = RGBColor(120, 120, 120)


def add_body_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_paragraph_indent(p)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT, ASCII_FONT, Pt(12))


def add_action(document: Document, number: str, title: str, body: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    set_keep_with_next(p)
    run = p.add_run(f"{number}. {title}")
    set_run_font(run, HEADING_FONT, ASCII_FONT, Pt(12))
    run.bold = True
    add_body_paragraph(document, body)


def add_section_title(document: Document, title: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_keep_with_next(p)
    run = p.add_run(title)
    set_run_font(run, HEADING_FONT, ASCII_FONT, Pt(12))
    run.bold = True


def add_image(document: Document, image: ExtractedImage, caption: str) -> None:
    image_path = Path(image.path)
    if not image_path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    width = Cm(12.8)
    if image.width and image.height and image.width < image.height:
        width = Cm(8.8)
    run.add_picture(str(image_path), width=width)

    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(8)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, BODY_FONT, ASCII_FONT, Pt(10.5))
    cap_run.font.color.rgb = RGBColor(90, 90, 90)


def add_toc_note(document: Document) -> None:
    document.add_heading("目录", level=1)
    add_body_paragraph(document, "目录可在 Word 中通过“引用—目录”自动生成；本稿已使用标准标题样式。")
    document.add_section(WD_SECTION.NEW_PAGE)


def add_preface(document: Document) -> None:
    document.add_heading("译制说明", level=1)
    add_body_paragraph(
        document,
        "本稿依据原书 PDF 文本进行中文译制，面向中国大陆举重训练、体能训练与专项力量训练读者。译文优先采用国内常用举重术语，并对编号动作、动作说明、页码来源和插图关系进行了中文书籍化重建。",
    )
    add_body_paragraph(
        document,
        "由于原书为训练动作资料汇编，部分页面以动作清单和示意图为主。本稿保留原书页码来源，并将嵌入图片作为正文插图处理，便于后续人工校对、补图注和终排。",
    )


def render_book_docx(
    output_path: Path,
    pages: list[RebuiltPage],
    images: list[ExtractedImage],
    title: str = "344 个中国举重运动员使用的奥运举重训练动作",
) -> None:
    document = Document()
    configure_document(document)

    add_title_page(document, title)
    add_preface(document)
    add_toc_note(document)

    by_page: dict[int, list[ExtractedImage]] = defaultdict(list)
    for image in images:
        by_page[image.page_number].append(image)

    for page in pages:
        add_source_page_marker(document, page.page_number)
        for paragraph in page.intro:
            add_body_paragraph(document, paragraph)
        if page.elements:
            for kind, value in page.elements:
                if kind == "section":
                    add_section_title(document, str(value))
                else:
                    add_action(document, value.number, value.title, value.body)
        else:
            for section in page.sections:
                add_section_title(document, section)
            for action in page.actions:
                add_action(document, action.number, action.title, action.body)
        for idx, image in enumerate(by_page.get(page.page_number, []), start=1):
            add_image(document, image, f"图 {page.page_number}-{idx} 原书第 {page.page_number} 页插图")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
