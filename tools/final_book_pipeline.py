from __future__ import annotations

import argparse
import html
import json
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

import fitz
from PIL import Image, ImageDraw, ImageFont, UnidentifiedImageError
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    from extract_pdf_images import ExtractedImage, extract_pdf_images
    from text_polisher import remove_ai_meta, strip_markdown_marks
    from terminology import enforce_terms
    from text_encoding_repair import repair_pdf_glyph_text
except ImportError:
    from .extract_pdf_images import ExtractedImage, extract_pdf_images
    from .text_polisher import remove_ai_meta, strip_markdown_marks
    from .terminology import enforce_terms
    from .text_encoding_repair import repair_pdf_glyph_text


BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"


@dataclass
class ParagraphBlock:
    text: str
    style: str = "body"


@dataclass
class TableBlock:
    caption: str | None
    rows: list[list[str]]


@dataclass
class ImageBlock:
    path: Path
    caption: str | None = None


Block = ParagraphBlock | TableBlock | ImageBlock


FIGURE_CAPTION_RE = re.compile(r"\s*Figure\s+((?:[A-Z]\.\d+)|(?:\d+(?:\.\d+)*))[a-z]?", re.IGNORECASE)
SOURCE_FIGURE_CAPTION_RE = re.compile(
    r"^\s*(?:Figure|Fig\.?|Figura|Рис\.?|Рисунок)\s+((?:[A-Z]\.\d+)|(?:\d+(?:\.\d+)*))[a-zа-я]?",
    re.IGNORECASE,
)
REFERENCE_HEADING_RE = re.compile(r"References\s+for\s+Chapter\s+(\d+)\.?", re.IGNORECASE)
RUSSIAN_REFERENCE_HEADING_RE = re.compile(
    r"^(?:\u041b\u0418\u0422\u0415\u0420\u0410\u0422\u0423\u0420\u0410|\u0421\u041f\u0418\u0421\u041e\u041a\s+\u041b\u0418\u0422\u0415\u0420\u0410\u0422\u0423\u0420\u042b)$",
    re.IGNORECASE,
)
CHINESE_FIGURE_CAPTION_PATTERN = r"图\s*(?:(?:[A-Z]\.\d+)|(?:\d+(?:\.\d+)?))[a-z]?(?:(?:[．、]\s*)|(?:\.\s+)|\s+).+"
CHINESE_TABLE_CAPTION_PATTERN = r"表\s*(?:(?:[A-Z]\.\d+)|(?:\d+(?:\.\d+)?))(?:(?:[．、]\s*)|(?:\.\s+)|\s+).+"
ALLOWED_SECTION_NUMBERS: set[tuple[int, ...]] = {
    (1, 1), (1, 1, 1), (1, 1, 2), (1, 1, 3),
    (1, 2), (1, 2, 1), (1, 2, 2), (1, 2, 3),
    (1, 3), (1, 3, 1), (1, 3, 2),
    (1, 4), (1, 4, 1), (1, 4, 2), (1, 4, 3), (1, 4, 4), (1, 4, 5),
    (2, 1), (2, 2), (2, 2, 1), (2, 2, 2), (2, 2, 3),
    (2, 3), (2, 3, 1), (2, 3, 2),
    (2, 4), (2, 4, 1), (2, 4, 2), (2, 4, 3),
    (2, 5), (2, 6), (2, 6, 1), (2, 6, 2),
    (3, 1), (3, 1, 1), (3, 1, 2), (3, 1, 3),
    (3, 2), (3, 2, 1), (3, 2, 2),
    (3, 3), (3, 3, 1), (3, 3, 2), (3, 3, 3),
    (4, 1), (4, 1, 1), (4, 1, 2), (4, 1, 3),
    (4, 2), (4, 2, 1), (4, 2, 2), (4, 2, 3), (4, 2, 4),
    (4, 3), (4, 3, 1), (4, 3, 2), (4, 3, 3), (4, 3, 4),
    (4, 4), (4, 4, 1), (4, 4, 2), (4, 4, 3),
    (4, 5), (4, 5, 1), (4, 5, 2),
    (4, 6), (4, 6, 1), (4, 6, 2),
    (5, 1), (5, 2), (5, 2, 1), (5, 2, 2),
    (5, 3), (5, 3, 1), (5, 3, 2), (5, 3, 3), (5, 3, 4),
    (5, 4), (5, 4, 1), (5, 4, 2), (5, 4, 3),
}
STANDARD_SECTION_TITLES: dict[tuple[int, ...], str] = {
    (1, 1): "运动训练与竞技准备的实质",
    (1, 1, 1): "目标、目的与训练指标",
    (1, 1, 2): "竞技准备的基本术语",
    (1, 1, 3): "训练方法",
    (1, 2): "训练与适应原则",
    (1, 2, 1): "训练负荷量与超负荷原则",
    (1, 2, 2): "训练负荷的专项性",
    (1, 2, 3): "适应",
    (1, 3): "超量恢复原则及其在实践中的应用",
    (1, 3, 1): "单次负荷后的超量恢复周期",
    (1, 3, 2): "训练课系列中多次负荷的累积效应",
    (1, 4): "运动训练的专项性原则",
    (1, 4, 1): "专项化",
    (1, 4, 2): "个体化",
    (1, 4, 3): "多样性",
    (1, 4, 4): "负荷交互作用",
    (1, 4, 5): "周期性训练设计",
    (2, 1): "训练效应：总体概述",
    (2, 2): "急性训练效应",
    (2, 2, 1): "通过专项指标评估急性训练效应",
    (2, 2, 2): "通过心理生理变量评估急性训练效应",
    (2, 2, 3): "急性训练效应的程序设计",
    (2, 3): "短时训练效应",
    (2, 3, 1): "短时训练效应的指标",
    (2, 3, 2): "短时训练效应的监控",
    (2, 4): "累积训练效应",
    (2, 4, 1): "生理变量的改善幅度",
    (2, 4, 2): "运动能力的提升",
    (2, 4, 3): "运动成绩的提升",
    (2, 5): "延迟训练效应",
    (2, 6): "残留训练效应",
    (2, 6, 1): "残留训练效应的基本概念与类型",
    (2, 6, 2): "影响短期残留训练效应的因素",
    (3, 1): "与遗传相关的可训练性决定因素",
    (3, 1, 1): "杰出的运动家族",
    (3, 1, 2): "身体与生理特征的遗传决定因素",
    (3, 1, 3): "累积训练效应的遗传决定因素",
    (3, 2): "可训练性与运动水平",
    (3, 2, 1): "可训练性的长期趋势",
    (3, 2, 2): "高反应者与低反应者",
    (3, 3): "可训练性与性别差异",
    (3, 3, 1): "最大运动表现的性别差异",
    (3, 3, 2): "运动适能生理决定因素的性别差异",
    (3, 3, 3): "训练反应中的性别差异",
    (4, 1): "训练周期化的基础",
    (4, 1, 1): "传统周期化的范畴",
    (4, 1, 2): "传统周期化的主要局限性",
    (4, 1, 3): "板块周期化：一种替代性训练理念",
    (4, 2): "训练课的特征与设计",
    (4, 2, 1): "基于负荷水平的训练课类型划分",
    (4, 2, 2): "高度集中工作负荷的训练课：关键训练课",
    (4, 2, 3): "训练课内的工作负荷集中：关键练习（任务）",
    (4, 2, 4): "不同工作负荷的兼容性",
    (4, 3): "小周期计划设计",
    (4, 3, 1): "不同类型小周期的特点与细节",
    (4, 3, 2): "单峰、双峰与三峰设计",
    (4, 3, 3): "针对不同训练模式的小周期",
    (4, 3, 4): "构建小周期计划的一般方法",
    (4, 4): "中周期计划设计",
    (4, 4, 1): "积累中周期",
    (4, 4, 2): "转化中周期",
    (4, 4, 3): "实现中周期",
    (4, 5): "年度周期设计",
    (4, 5, 1): "年度计划的基础",
    (4, 5, 2): "年度备战计划示例",
    (4, 6): "针对目标赛事的赛前最终阶段备战",
    (4, 6, 1): "影响赛前最终阶段备战效果的因素",
    (4, 6, 2): "赛前最终阶段备战的内容与细节",
    (5, 1): "运动员竞技准备的通用模型",
    (5, 2): "顶级竞技表现模型",
    (5, 2, 1): "个人项目",
    (5, 2, 2): "集体项目",
    (5, 3): "专项能力模型",
    (5, 3, 1): "专项能力的通用因素",
    (5, 3, 2): "体型与身体成分",
    (5, 3, 3): "生理能力",
    (5, 3, 4): "专项运动能力",
    (5, 4): "训练计划模型",
    (5, 4, 1): "结构模型",
    (5, 4, 2): "训练内容模型",
    (5, 4, 3): "训练负荷的模型特征",
}

CHINESE_WEIGHTLIFTING_CHAPTER_TITLES: dict[int, str] = {
    1: "第一章 中国举重的历史与发展",
    2: "第二章 中国举重理念：三原则与五字诀",
    3: "第三章 举重中的爆发力、力量与速度",
    4: "第四章 举重执教方法",
    5: "第五章 举重训练原则",
    6: "第六章 运动训练",
    7: "第七章 技术训练",
    8: "第八章 心理训练",
    9: "第九章 战术训练",
    10: "第十章 恢复训练",
    11: "第十一章 智力训练",
    12: "第十二章 意志品质训练",
    13: "第十三章 训练负荷",
    14: "第十四章 力量训练理论与方法",
    15: "第十五章 训练计划与训练日记",
    16: "第十六章 运动员评估",
    17: "第十七章 青少年训练",
    18: "第十八章 女子举重",
    19: "第十九章 举重运动员选材方法",
    20: "第二十章 运动损伤",
}

SOCCER_CHAPTER_TITLES: dict[int, str] = {
    1: "第一章 从业者的角色",
    2: "第二章 从青训到职业",
    3: "第三章 需求分析与测试",
    4: "第四章 力量、爆发力与损伤预防",
    5: "第五章 体能训练",
    6: "第六章 球员监控与实操应用",
    7: "第七章 可穿戴技术",
    8: "第八章 恢复与营养",
    9: "第九章 重返赛场",
    10: "第十章 周期化安排",
    11: "第十一章 教练与工作人员整合",
}

SOCCER_CHAPTER_ALIASES: dict[str, int] = {
    "从业者的角色": 1,
    "从青训到职业": 2,
    "从青训学院到职业赛场": 2,
    "需求分析与测试": 3,
    "力量、爆发力与损伤预防": 4,
    "力量、爆发力与损伤": 4,
    "体能训练": 5,
    "球员监控与实操应用": 6,
    "球员监控与实践应用": 6,
    "球员监控与": 6,
    "实操应用": 6,
    "实践应用": 6,
    "可穿戴技术": 7,
    "恢复与营养": 8,
    "营养与恢复": 8,
    "重返赛场": 9,
    "周期化安排": 10,
    "教练与工作人员整合": 11,
    "教练与工作人员的整合": 11,
    "教练与团队整合": 11,
}

SOCCER_CHAPTER_START_PAGES: dict[int, int] = {
    24: 1,
    44: 2,
    80: 3,
    106: 4,
    131: 5,
    162: 6,
    188: 7,
    212: 8,
    246: 9,
    282: 10,
    312: 11,
}

SOCCER_GENERIC_BARE_HEADINGS = {
    "体能训练",
    "恢复与营养",
    "营养与恢复",
    "周期化安排",
    "球员监控与",
    "实操应用",
    "实践应用",
}

FRANK_CHAPTER_TITLES: dict[int, str] = {
    1: "第1章 工作部件",
    2: "第2章 儿童生长发育中的结构变化",
    3: "第3章 基础力学",
    4: "第4章 营养",
    5: "第5章 氧气运输系统",
    6: "第6章 工作中的肌肉",
    7: "第7章 体液系统",
    8: "第8章 激素",
    9: "第9章 成长中儿童的生理差异",
    10: "第10章 身心合一",
    11: "第11章 知觉-运动学习",
    12: "第12章 成长中儿童的心理变化",
    13: "第13章 体能",
    14: "第14章 力量发展的理论与实践",
    15: "第15章 速度发展的理论与实践",
    16: "第16章 耐力发展的理论与实践",
    17: "第17章 柔韧性发展的理论与实践",
    18: "第18章 运动评价",
    19: "第19章 年度周期化",
    20: "第20章 周期化的变化形式",
    21: "第21章 训练单元、小周期、中周期与大周期",
    22: "第22章 负荷适应",
    23: "第23章 训练与过度训练",
    24: "第24章 比赛期",
    25: "第25章 教练工作：原则与实践",
}

FRANK_CHAPTER_ALIASES: dict[str, int] = {
    "工作部件": 1,
    "儿童生长发育中的结构变化": 2,
    "成长中儿童的结构性变化": 2,
    "成长中儿童的结构变化": 2,
    "基础力学": 3,
    "营养": 4,
    "营养学": 4,
    "氧气运输系统": 5,
    "工作中的肌肉": 6,
    "工作肌肉": 6,
    "体液系统": 7,
    "激素": 8,
    "成长中儿童的生理差异": 9,
    "生长发育中儿童的生理差异": 9,
    "身心合一": 10,
    "志在征服，体在征服": 10,
    "知觉-运动学习": 11,
    "成长中儿童的心理变化": 12,
    "心理变化与成长中的儿童": 12,
    "体能": 13,
    "力量发展的理论与实践": 14,
    "速度发展的理论与实践": 15,
    "耐力发展的理论与实践": 16,
    "柔韧性发展的理论与实践": 17,
    "灵活性发展的理论与实践": 17,
    "活动度发展的理论与实践": 17,
    "运动评价": 18,
    "运动训练中的评价": 18,
    "年度周期化": 19,
    "周期化的变化形式": 20,
    "训练单元、小周期、中周期与大周期": 21,
    "训练课、小周期、中周期与大周期": 21,
    "训练课、小周期、 中周期与大周期": 21,
    "负荷适应": 22,
    "对负荷的适应": 22,
    "训练与过度训练": 23,
    "比赛期": 24,
    "教练工作：原则与实践": 25,
    "教练工作: 原则与实践": 25,
}

GOMES_CHAPTER_TITLES: dict[int, str] = {
    1: "第1章 竞技准备的科学原则",
    2: "第2章 运动竞赛体系",
    3: "第3章 竞技准备的手段与方法",
    4: "第4章 训练负荷",
    5: "第5章 身体能力的训练与完善",
    6: "第6章 运动训练的结构化与周期化",
    7: "第7章 运动项目中的周期化模型",
    8: "第8章 儿童与青少年运动训练计划",
    9: "第9章 运动训练计划",
}

GOMES_CHAPTER_ALIASES: dict[str, int] = {
    "竞技准备的科学原则": 1,
    "竞技准备的科学原理": 1,
    "运动训练的科学原则": 1,
    "科学原则": 1,
    "运动竞赛体系": 2,
    "竞技准备的手段与方法": 3,
    "训练负荷": 4,
    "身体能力的训练与完善": 5,
    "身体能力的训练与提高": 5,
    "体能能力的训练与完善": 5,
    "体能能力的训练与提高": 5,
    "运动训练的结构化与周期化": 6,
    "运动训练结构化与周期化": 6,
    "运动项目中的周期化模型": 7,
    "儿童与青少年运动训练计划": 8,
    "儿童与青少年时期的运动训练": 8,
    "运动训练计划": 9,
}

GOMES_CHAPTER_START_PAGES: dict[int, int] = {
    18: 1,
    34: 2,
    51: 3,
    67: 4,
    86: 5,
    144: 6,
    204: 7,
    227: 8,
    248: 9,
}

GILBERT_CHAPTER_TITLES: dict[int, str] = {
    1: "第一章 明确使命与核心价值观",
    2: "第二章 将价值观融入执教理念",
    3: "第三章 设定目标成果",
    4: "第四章 建立信任与凝聚力",
    5: "第五章 培养运动天赋与技能",
    6: "第六章 优化运动员的学习",
    7: "第七章 设计有效的训练环境",
    8: "第八章 确保赛前准备状态",
    9: "第九章 在比赛日有效执教",
    10: "第十章 设计与实施评估系统",
    11: "第十一章 识别并发挥优势",
    12: "第十二章 缩小表现差距",
    13: "第十三章 协作与学习",
    14: "第十四章 充电与激发",
}

GILBERT_CHAPTER_ALIASES: dict[str, int] = {
    "明确使命与核心价值观": 1,
    "将价值观融入执教理念": 2,
    "将价值观与执教理念相联结": 2,
    "连接价值观与理念": 2,
    "设定目标成果": 3,
    "建立信任与凝聚力": 4,
    "培养运动天赋与技能": 5,
    "发展运动天赋与技能": 5,
    "优化运动员的学习": 6,
    "优化运动员的学习过程": 6,
    "设计有效的训练环境": 7,
    "确保赛前准备状态": 8,
    "在比赛日有效执教": 9,
    "比赛日高效执教": 9,
    "设计与实施评估系统": 10,
    "设计并实施评估系统": 10,
    "识别并发挥优势": 11,
    "识别并强化优势": 11,
    "缩小表现差距": 12,
    "协作与学习": 13,
    "充电与激发": 14,
    "充电与点燃": 14,
}

GILBERT_CHAPTER_START_PAGES: dict[int, int] = {
    21: 1,
    56: 2,
    89: 3,
    121: 4,
    165: 5,
    190: 6,
    214: 7,
    250: 8,
    291: 9,
    336: 10,
    413: 11,
    446: 12,
    482: 13,
    510: 14,
}

GILBERT_PART_TITLES: dict[int, str] = {
    20: "第一部分 赛季前：愿景",
    164: "第二部分 赛季中：执行",
    335: "第三部分 赛季结束：评估",
    445: "第四部分 休赛期：提升",
}

GILBERT_SKIP_PAGES = {6, 10, 11, 12}


GORDON_CHAPTER_TITLES: dict[int, str] = {
    1: "第一章 执教过程",
    2: "第二章 领导力与教练员—运动员关系",
    3: "第三章 技能习得与学习",
    4: "第四章 目标与目标设定",
    5: "第五章 动机",
    6: "第六章 焦虑、压力与表现",
    7: "第七章 训练变量与组成要素",
    8: "第八章 训练理论与模型",
    9: "第九章 训练计划与结构安排",
    10: "第十章 耐力训练",
    11: "第十一章 力量与爆发力训练",
    12: "第十二章 柔韧性训练",
    13: "第十三章 速度、灵敏与快速能力",
    14: "第十四章 运动员选材",
    15: "第十五章 长期运动员发展",
}

GORDON_CHAPTER_ALIASES: dict[str, int] = {
    "执教过程": 1,
    "教练过程": 1,
    "执教流程": 1,
    "领导力与教练员—运动员关系": 2,
    "领导力与教练员-运动员关系": 2,
    "领导力与教练员运动员关系": 2,
    "技能习得与学习": 3,
    "技能获得与学习": 3,
    "目标与目标设定": 4,
    "目标与目标设置": 4,
    "动机": 5,
    "焦虑、压力与表现": 6,
    "焦虑压力与表现": 6,
    "训练变量与组成要素": 7,
    "训练变量与训练组成要素": 7,
    "训练变量和组成要素": 7,
    "训练理论与模型": 8,
    "训练计划与结构安排": 9,
    "训练计划与结构": 9,
    "训练规划与结构安排": 9,
    "耐力训练": 10,
    "力量与爆发力训练": 11,
    "力量和爆发力训练": 11,
    "柔韧性训练": 12,
    "速度、灵敏与快速能力": 13,
    "速度灵敏与快速能力": 13,
    "速度、敏捷与快速能力": 13,
    "运动员选材": 14,
    "人才识别": 14,
    "天赋识别": 14,
    "长期运动员发展": 15,
    "长期运动发展": 15,
}

GORDON_CHAPTER_START_PAGES: dict[int, int] = {
    13: 1,
    24: 2,
    36: 3,
    47: 4,
    57: 5,
    66: 6,
    85: 7,
    99: 8,
    121: 9,
    149: 10,
    172: 11,
    187: 12,
    195: 13,
    213: 14,
    230: 15,
}

GORDON_PART_TITLES: dict[int, str] = {
    11: "第一部分 理解教练员",
    45: "第二部分 运动表现心理学",
    83: "第三部分 训练运动员",
    211: "第四部分 运动员监控与评估",
}

GORDON_SKIP_PAGES = {1, 3, 6, 10, 12}


LTAD_CHAPTER_TITLES: dict[int, str] = {
    1: "第一章 长期运动员发展模型",
    2: "第二章 身体素养",
    3: "第三章 运动专项化",
    4: "第四章 可训练性",
    5: "第五章 智力、情绪与道德发展",
    6: "第六章 体育系统协同与整合",
    7: "第七章 长期运动员发展的阶段",
}

LTAD_CHAPTER_ALIASES: dict[str, int] = {
    "长期运动员发展模型": 1,
    "长期运动发展模型": 1,
    "LTAD模型": 1,
    "身体素养": 2,
    "体育素养": 2,
    "体能素养": 2,
    "运动专项化": 3,
    "运动专门化": 3,
    "体育专项化": 3,
    "可训练性": 4,
    "训练适应性": 4,
    "智力、情绪与道德发展": 5,
    "智力情绪与道德发展": 5,
    "智力、情感与道德发展": 5,
    "智力情感与道德发展": 5,
    "认知、情绪与道德发展": 5,
    "体育系统协同与整合": 6,
    "运动系统协同与整合": 6,
    "体育系统对齐与整合": 6,
    "长期运动员发展的阶段": 7,
    "长期运动员发展阶段": 7,
    "长期运动发展阶段": 7,
}

LTAD_CHAPTER_START_PAGES: dict[int, int] = {
    19: 1,
    67: 2,
    115: 3,
    155: 4,
    199: 5,
    247: 6,
    287: 7,
}

LTAD_FRONT_TITLES: dict[int, str] = {
    17: "前言",
}

LTAD_SKIP_PAGES = {1, 2, 3, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16}


STRENGTH_CHAPTER_TITLES: dict[int, str] = {
    1: "第一章 训练理论的基本概念",
    2: "第二章 任务专项力量",
    3: "第三章 运动员专项力量",
    4: "第四章 训练强度",
    5: "第五章 力量训练中的时机安排",
    6: "第六章 力量训练所用练习",
    7: "第七章 力量房中的速度",
    8: "第八章 损伤预防",
    9: "第九章 过度延伸、过度训练与恢复",
    10: "第十章 力量房中的运动员监控",
    11: "第十一章 目标导向力量训练",
    12: "第十二章 女性力量训练",
    13: "第十三章 青少年运动员力量训练",
    14: "第十四章 老年运动员力量训练",
}

STRENGTH_CHAPTER_ALIASES: dict[str, int] = {
    "训练理论的基本概念": 1,
    "训练理论基本概念": 1,
    "任务专项力量": 2,
    "任务特异性力量": 2,
    "运动员专项力量": 3,
    "运动员特异性力量": 3,
    "训练强度": 4,
    "力量训练中的时机安排": 5,
    "力量训练中的时间安排": 5,
    "力量训练所用练习": 6,
    "用于力量训练的练习": 6,
    "力量房中的速度": 7,
    "举重房中的速度": 7,
    "损伤预防": 8,
    "伤病预防": 8,
    "过度延伸、过度训练与恢复": 9,
    "过度训练与恢复": 9,
    "力量房中的运动员监控": 10,
    "目标导向力量训练": 11,
    "目标专项力量训练": 11,
    "女性力量训练": 12,
    "青少年运动员力量训练": 13,
    "老年运动员力量训练": 14,
}

STRENGTH_CHAPTER_START_PAGES: dict[int, int] = {
    21: 1,
    33: 2,
    59: 3,
    79: 4,
    97: 5,
    117: 6,
    143: 7,
    177: 8,
    195: 9,
    213: 10,
    231: 11,
    253: 12,
    275: 13,
    299: 14,
}

STRENGTH_PART_TITLES: dict[int, str] = {
    19: "第一部分 力量训练基础",
    77: "第二部分 力量训练的概念",
    251: "第三部分 特殊人群力量训练",
}

STRENGTH_FRONT_TITLES: dict[int, str] = {
    11: "序言",
    13: "前言",
    15: "致谢",
    17: "符号与缩略语",
}

STRENGTH_SKIP_PAGES = {1, 2, 3, 6, 7, 8, 9, 10, 12, 16, 18, 346}

PLATONOV_PART_TITLES: dict[int, str] = {
    1: "第一部分 运动训练分期理论的历史、总体结构与基本原则",
    2: "第二部分 运动训练分期理论的一般基础",
    3: "第三部分 训练过程的宏观结构与运动员选材、定向的分阶段体系",
    4: "第四部分 训练过程的微观与中观结构",
    5: "第五部分 运动员训练过程的宏观结构",
    6: "第六部分 训练与竞赛外部因素在运动员准备体系中的应用",
}

PLATONOV_CHAPTER_TITLES: dict[int, str] = {
    1: "第1章 运动训练分期——发展历史",
    2: "第2章 对分期理论的批判、替代性概念与创新性方法",
    3: "第3章 总体结构、内容及概念术语体系",
    4: "第4章 分期体系中的专项性原则与一般教学原则",
    5: "第5章 运动中的适应",
    6: "第6章 个体发育与适应过程",
    7: "第7章 骨骼肌：结构、功能与适应",
    8: "第8章 肌肉活动的能量供应系统",
    9: "第9章 负荷、疲劳、恢复、超量恢复与训练的长期效应",
    10: "第10章 随意运动控制的基础",
    11: "第11章 多年训练：知识体系的形成",
    12: "第12章 多年训练的规律与特征",
    13: "第13章 多年训练分期的现代体系",
    14: "第14章 多年训练准备体系中的选材与定向",
    15: "第15章 运动训练的手段",
    16: "第16章 热身及热身程序的制定",
    17: "第17章 训练课及其程序制定",
    18: "第18章 小周期及其程序制定",
    19: "第19章 中周期及其程序制定",
    20: "第20章 年度训练分期的基础",
    21: "第21章 东欧关于年度训练分期的经验",
    22: "第22章 美国与澳大利亚关于年度训练分期的经验",
    23: "第23章 年度训练分期的现代模式",
    24: "第24章 赛前直接准备",
    25: "第25章 中高海拔地区训练与人工低氧在运动员准备体系中的应用",
    26: "第26章 昼夜节律紊乱条件下的运动员训练",
    27: "第27章 恢复手段与功能能力刺激方法",
    28: "第28章 营养学与药理学辅助",
}

PLATONOV_MOTOR_PART_TITLES: dict[int, str] = {
    1: "第一部分 运动员身体准备基础",
    2: "第二部分 运动素质及其发展",
    3: "第三部分 休息优化、恢复与工作能力促进手段及膳食保障",
    4: "第四部分 过度训练与运动损伤",
}

PLATONOV_MOTOR_PART_START_PAGES: dict[int, int] = {
    13: 1,
    267: 2,
    517: 3,
    565: 4,
}

PLATONOV_MOTOR_CHAPTER_TITLES: dict[int, str] = {
    1: "第一章 运动员的运动素质与身体准备",
    2: "第二章 适应理论与功能系统理论在运动员准备知识体系发展中的作用",
    3: "第三章 骨骼-肌肉系统",
    4: "第四章 氧运输系统",
    5: "第五章 肌肉活动的能量供应",
    6: "第六章 运动员运动活动的神经调节与控制",
    7: "第七章 运动员身体准备中的心理调节",
    8: "第八章 运动员的年龄发展与身体准备",
    9: "第九章 身体准备方法差异的依据",
    10: "第十章 运动员身体准备系统中的负荷",
    11: "第十一章 准备活动：意义、一般与专项构建基础",
    12: "第十二章 力量及其发展方法",
    13: "第十三章 灵敏性与协调能力及其发展方法",
    14: "第十四章 速度能力及其发展方法",
    15: "第十五章 柔韧性及其发展方法",
    16: "第十六章 耐力及其发展方法",
    17: "第十七章 恢复与工作能力促进手段",
    18: "第十八章 膳食营养保障",
    19: "第十九章 现代运动中的过度训练问题",
    20: "第二十章 运动损伤与疾病",
}

PLATONOV_MOTOR_CHAPTER_START_PAGES: dict[int, int] = {
    14: 1,
    27: 2,
    49: 3,
    74: 4,
    90: 5,
    126: 6,
    153: 7,
    163: 8,
    189: 9,
    202: 10,
    259: 11,
    268: 12,
    358: 13,
    390: 14,
    416: 15,
    440: 16,
    518: 17,
    535: 18,
    566: 19,
    585: 20,
}

PLATONOV_MOTOR_SKIP_PAGES = {1, 2, 3, 4, 5, 6, 7, 12}


BONDARCHUK_CHAPTER_TITLES: dict[int, str] = {
    1: "第一章 训练迁移问题研究的历史简述",
    2: "第二章 竞技状态形成周期与训练迁移",
    3: "第三章 使用不同种类练习时运动能力的迁移",
    4: "第四章 运动技能的迁移",
    5: "第五章 各训练课、课中部分及练习之间的相互关系",
    6: "第六章 运动训练中单个训练课、小周期、中周期和大周期训练负荷的标准化",
    7: "第七章 机体系统对促进运动成绩提高的适应性改造的保护",
}

BONDARCHUK_CHAPTER_START_PAGES: dict[int, int] = {
    5: 1,
    22: 2,
    56: 3,
    124: 4,
    152: 5,
    176: 6,
    227: 7,
}

BONDARCHUK_SKIP_PAGES = {1, 270, 271}


SCIENCE_SOCCER_SECTION_TITLES: dict[int, str] = {
    24: "A部分 生物科学",
    132: "B部分 社会与行为科学",
    220: "C部分 运动医学与生物力学",
    274: "D部分 表现分析与监控",
    350: "E部分 人才识别、生长与发展",
    418: "F部分 俱乐部中的关键组织角色",
}

SCIENCE_SOCCER_CHAPTER_TITLES: dict[int, str] = {
    1: "第一章 身体准备",
    2: "第二章 抗阻训练",
    3: "第三章 有氧与无氧训练",
    4: "第四章 热环境下的足球：表现与缓解策略",
    5: "第五章 比赛与训练营养",
    6: "第六章 恢复策略",
    7: "第七章 球员心理特征",
    8: "第八章 预判与决策",
    9: "第九章 技能习得：球员发展路径与有效练习",
    10: "第十章 球员识别与发展中的社会学影响",
    11: "第十一章 球员身心健康与职业转型",
    12: "第十二章 发展适应性执教专长",
    13: "第十三章 损伤流行病学、监控与预防",
    14: "第十四章 传染性疾病",
    15: "第十五章 生物力学评估",
    16: "第十六章 比赛中体能表现分析",
    17: "第十七章 技战术比赛分析",
    18: "第十八章 训练监控",
    19: "第十九章 训练与比赛负荷数据的应用",
    20: "第二十章 生长与成熟",
    21: "第二十一章 天赋突出，还是发育领先？如何改进球员评价",
    22: "第二十二章 人才识别与人才促进",
    23: "第二十三章 现代球探与招募方法",
    24: "第二十四章 运动科学总监或高表现总监的工作",
    25: "第二十五章 体育总监的工作",
}

SCIENCE_SOCCER_CHAPTER_START_PAGES: dict[int, int] = {
    26: 1,
    38: 2,
    57: 3,
    75: 4,
    90: 5,
    113: 6,
    134: 7,
    147: 8,
    164: 9,
    178: 10,
    191: 11,
    206: 12,
    222: 13,
    246: 14,
    261: 15,
    276: 16,
    296: 17,
    315: 18,
    332: 19,
    352: 20,
    369: 21,
    386: 22,
    405: 23,
    420: 24,
    437: 25,
}

SCIENCE_SOCCER_SKIP_PAGES = {1, 3, 6, 7, 8, 451}

SOVIET_WEIGHTLIFTING_TITLES: dict[int, str] = {
    1: "本书适合哪些读者",
    2: "读者推荐语",
    3: "致谢与前言",
    4: "什么是苏联举重体系",
    5: "苏联举重体系的奠基人",
    6: "苏联举重体系的架构",
    7: "苏联举重体系的哲学基础",
    8: "提升竞技运动表现的策略与方法",
    9: "优化训练负荷的手段与方法",
    10: "发展力量与爆发力练习的动力学和运动学结构",
    11: "制定训练计划",
    12: "案例研究：俄罗斯、中国与保加利亚女子举重运动员的差异",
    13: "示例训练计划",
    14: "十二周训练计划",
    15: "每周三次训练计划",
    16: "每周四次训练计划",
    17: "每周五次训练计划",
    18: "每周六次训练计划",
    19: "参考文献",
    20: "作者简介",
    21: "关于 Risto Sports",
}

SOVIET_WEIGHTLIFTING_START_PAGES: dict[int, int] = {
    10: 1,
    11: 2,
    13: 3,
    14: 4,
    16: 5,
    19: 6,
    21: 7,
    23: 8,
    29: 9,
    36: 10,
    45: 11,
    61: 12,
    64: 13,
    67: 14,
    70: 15,
    75: 16,
    82: 17,
    89: 18,
    97: 19,
    100: 20,
    104: 21,
}

SOVIET_WEIGHTLIFTING_SKIP_PAGES = {6, 7, 8, 9, 107, 108}


def set_run_font(run, east_asian: str = BODY_FONT, latin: str = LATIN_FONT, size: Pt | None = None) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asian)
    if size is not None:
        run.font.size = size


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def xml_safe_text(text: str | None) -> str:
    """Remove characters that cannot be stored in DOCX XML text nodes."""
    if not text:
        return ""
    safe_chars = []
    for char in str(text):
        codepoint = ord(char)
        if char in {"\t", "\n", "\r"}:
            safe_chars.append(char)
        elif 0x20 <= codepoint <= 0xD7FF or 0xE000 <= codepoint <= 0xFFFD or 0x10000 <= codepoint <= 0x10FFFF:
            safe_chars.append(char)
    return "".join(safe_chars)


def paragraph_display_text(text: str | None) -> str:
    text = xml_safe_text(text)
    text = re.sub(r"\s*\|\s*", "；", text)
    text = re.sub(r"；{2,}", "；", text)
    return text.strip()


def clean_text(text: str) -> str:
    text = xml_safe_text(text)
    text = remove_ai_meta(text)
    text = strip_markdown_marks(text)
    text = enforce_terms(text)
    text = re.sub(r"(?mi)^.*第\s*X\s*(?:章|部分).*$", "", text)
    text = re.sub(r"(?m)^.*原书第\s*\d+\s*页.*$", "", text)
    text = re.sub(r"原书第\s*\d+\s*页", "", text)
    text = re.sub(r"(?m)^.*(?:原文此处|OCR).*?(?:乱码|不完整|截断|无法识别|无法翻译).*?$", "", text)
    text = re.sub(r"(?m)^.*译者根据上下文.*?$", "", text)
    text = re.sub(r"(?m)^.*(?:版权信息|版权所有|版权归).*?(?:БИБКОМ|Книга-Cервис|ООО).*$", "", text)
    text = re.sub(
        r"从\s*W\.\s*Gilbert\s*著《[^。\n]+?Human Kinetics[）)](?:中摘录|开始)。?",
        "资料来源：W. Gilbert, Coaching Better Every Season: A Year-Round System for Athlete Development and Program Success (Champaign, IL: Human Kinetics, 2017).",
        text,
    )
    text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    text = re.sub(r"(?m)^\s*[<>]{2,}\s*$", "", text)
    text = re.sub(r"[<>]{4,}", "", text)
    text = re.sub(r"(?mi)^\s*(?:plaintext|markdown|text)\s*$", "", text)
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
    text = re.sub(r"\S+?\.qxd:[^\n]*?(?:Page\s*\d+|第\s*\d+\s*页)\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"(?m)^.*853\w*_[^\n]*(?:Sports Training Principles|运动训练原则)?[^\n]*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"(?m)^\s*(?:SPORTS TRAINING PRINCIPLES|运动训练原则|运动训练原理|trahea)\s*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^#\s*第\s*\d+\s*页\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"(?m)^#{1,6}\s*", "", text)
    text = re.sub(r"(?m)^\s*第\s*\d+\s*页\s*$", "", text)
    text = re.sub(r"(?mi)^\s*t\s*\n\s*r\s*\n\s*a\s*\n\s*c\s*\n\s*h\s*\n\s*e\s*\n\s*a\s*$", "气管", text)
    text = re.sub(r"(?m)^\s*(?:[ivxlcdm]+|\d+)\s*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = remove_leading_chapter_listing(text)
    return text.strip()


def is_short_chapter_listing_line(text: str) -> bool:
    stripped = text.strip()
    if len(stripped) > 90:
        return False
    return bool(re.match(r"^第\s*[一二三四五六七八九十\d]+\s*章\b", stripped))


def remove_leading_chapter_listing(text: str) -> str:
    """Remove part title pages' chapter lists while keeping the part title and body."""
    lines = text.splitlines()
    nonempty = [(index, line.strip()) for index, line in enumerate(lines) if line.strip()]
    if not nonempty:
        return text
    first_window = nonempty[:18]
    listing_count = sum(1 for _index, line in first_window if is_short_chapter_listing_line(line))
    if listing_count < 3:
        return text
    removable_indices = {
        index
        for index, line in first_window
        if is_short_chapter_listing_line(line)
    }
    cleaned = [line for index, line in enumerate(lines) if index not in removable_indices]
    return "\n".join(cleaned)


CHUNK_MARKER_RE = re.compile(r"<!--\s*page\s+(\d+),\s*chunk\s+(\d+)/(\d+)\s*-->", re.IGNORECASE)


def remove_translation_meta_lines(text: str) -> str:
    """Strip model-side explanations while preserving translated content."""
    removed_meta = []
    skip_rest = False
    for line in text.splitlines():
        stripped = line.strip()
        compact = re.sub(r"\s+", "", stripped)
        if not stripped:
            removed_meta.append(line)
            continue
        if stripped in {"---", "——"}:
            continue
        if re.match(r"^\*{0,2}说明\*{0,2}[:：]?$", stripped):
            skip_rest = True
            continue
        if skip_rest:
            continue
        if any(
            token in stripped
            for token in [
                "抱歉",
                "请提供",
                "请粘贴",
                "请将第",
                "请将需要翻译",
                "如果您能提供",
                "如果您需要",
                "若需",
                "以下是您提供",
                "好的，这是",
                "示例翻译",
                "无法识别",
                "无法翻译",
                "无法确定",
                "译者根据上下文",
                "此处原文可能",
                "原文可能因OCR",
                "OCR修复",
                "术语统一",
                "内容忠实",
                "语言风格",
                "版权信息",
                "版权所有",
                "版权归",
            ]
        ):
            continue
        if re.search(r"原书第\s*\d+\s*页", stripped):
            continue
        if re.search(r"(?:原文此处|OCR).*?(?:乱码|不完整|截断|无法识别|无法翻译)", stripped):
            continue
        removed_meta.append(line)
    return "\n".join(removed_meta).strip()


def looks_like_synthetic_outline_chunk(body: str, chunk_index: int, chunk_total: int) -> bool:
    if chunk_total <= 1 or chunk_index != 1:
        return False
    heading_count = len(re.findall(r"(?m)^\s*#{1,4}\s+", body))
    if "参考文献" in body:
        return True
    if re.search(r"第\s*X\s*章", body, flags=re.IGNORECASE):
        return True
    generic_title_tokens = [
        "运动训练分期理论的基本概念",
        "训练分期理论的基本概念",
        "运动训练分期理论的核心概念",
        "训练分期理论的应用",
        "年度训练计划的结构",
        "年度训练计划的制定",
        "未来发展方向",
        "本章小结",
    ]
    if heading_count >= 4 and any(token in body for token in generic_title_tokens):
        return True
    if (
        heading_count >= 5
        and len(body) > 1600
        and ("运动训练分期" in body or "训练分期" in body)
        and re.search(r"(?m)^\s*#{1,4}\s*(?:\d+\.\d+|第\s*\d+\s*页)", body)
    ):
        return True
    return False


def filter_translation_artifact_chunks(raw_text: str) -> str:
    matches = list(CHUNK_MARKER_RE.finditer(raw_text))
    if not matches:
        return remove_translation_meta_lines(raw_text)
    output = [remove_translation_meta_lines(raw_text[: matches[0].start()])]
    for index, match in enumerate(matches):
        chunk_start = match.end()
        chunk_end = matches[index + 1].start() if index + 1 < len(matches) else len(raw_text)
        body = remove_translation_meta_lines(raw_text[chunk_start:chunk_end])
        try:
            chunk_index = int(match.group(2))
            chunk_total = int(match.group(3))
        except ValueError:
            chunk_index = 1
            chunk_total = 1
        if not body.strip():
            continue
        if looks_like_synthetic_outline_chunk(body, chunk_index, chunk_total):
            continue
        output.append(body.strip())
    return "\n\n".join(part for part in output if part and part.strip()).strip()


def source_text_by_page(pdf_path: Path) -> dict[int, str]:
    pages: dict[int, str] = {}
    with fitz.open(pdf_path) as doc:
        for page_index, page in enumerate(doc, start=1):
            pages[page_index] = repair_pdf_glyph_text(page.get_text("text"))
    return pages


def is_image_only_pdf(source_pages: dict[int, str]) -> bool:
    if not source_pages:
        return False
    nonempty = sum(1 for text in source_pages.values() if len(text.strip()) > 30)
    return nonempty <= max(2, len(source_pages) * 0.03)


def trim_low_text_page_hallucination(text: str, source_text: str) -> str:
    source_compact = " ".join(source_text.split())
    if len(source_compact) > 160:
        return text
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return text
    kept: list[str] = []
    source_is_figure = bool(SOURCE_FIGURE_CAPTION_RE.match(source_compact))
    source_is_credit = source_compact.startswith(("Reprinted", "Adapted", "©", "Copyright"))
    for line in lines:
        if re.match(rf"^{CHINESE_FIGURE_CAPTION_PATTERN}$", line):
            kept.append(line)
            continue
        if re.match(r"^图\s*(?:[A-Z]\.\d+|[A-Z]?\d+(?:\.\d+)?)[a-z]?\s*.+", line):
            kept.append(line)
            continue
        if source_is_credit and (line.startswith(("经许可转载", "改编自", "©")) or "Human Kinetics" in line):
            kept.append(line)
            continue
        if source_is_figure and len(kept) >= 1:
            break
    if kept and len("".join(lines)) > max(500, len("".join(kept)) * 3):
        return "\n".join(kept)
    if source_is_credit and kept:
        return "\n".join(kept[:1])
    return text


def trim_generated_figure_artifacts(text: str, source_text: str) -> str:
    source_compact = " ".join(source_text.split())
    source_starts_with_figure = bool(SOURCE_FIGURE_CAPTION_RE.match(source_compact))
    source_starts_with_table = bool(re.match(r"^Table\s+\d", source_compact, flags=re.I))
    if source_starts_with_figure or source_starts_with_table:
        return text
    lines = text.splitlines()
    for index, line in enumerate(lines):
        stripped = line.strip()
        if re.match(rf"^(?:{CHINESE_FIGURE_CAPTION_PATTERN}|{CHINESE_TABLE_CAPTION_PATTERN})$", stripped):
            tail = lines[index + 1 :]
            has_table = any("|" in item for item in tail[:6])
            near_end = index >= max(0, len(lines) - 5)
            if has_table or near_end:
                return "\n".join(lines[:index]).strip()
    marker = re.search(
        rf"\n\s*---\s*\n(?=(?:\*\*)?\s*(?:{CHINESE_FIGURE_CAPTION_PATTERN}|{CHINESE_TABLE_CAPTION_PATTERN}))",
        text,
    )
    if marker:
        return text[: marker.start()].strip()
    marker = re.search(
        r"\n\s*---\s*\n(?=.*\|.*\|)",
        text,
        flags=re.DOTALL,
    )
    if marker:
        tail = text[marker.end() :]
        if re.search(rf"(?:{CHINESE_FIGURE_CAPTION_PATTERN}|{CHINESE_TABLE_CAPTION_PATTERN})", tail):
            return text[: marker.start()].strip()
    return text


def trim_low_text_figure_page(text: str, source_text: str) -> str:
    source_compact = " ".join(source_text.split())
    if len(source_compact) > 300:
        return text
    source_upper = source_compact.upper()
    has_figure_signal = (
        "\u0420\u0418\u0421\u0423\u041d\u041e\u041a" in source_upper
        or SOURCE_FIGURE_CAPTION_RE.search(source_compact) is not None
    )
    if not has_figure_signal:
        return text
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    kept: list[str] = []
    for line in lines:
        if re.match(r"^第\s*[一二三四五六七八九十\d]+\s*章", line):
            if kept:
                break
            continue
        if re.match(r"^第\s*[一二三四五六七八九十\d]+\s*部分", line):
            continue
        if re.match(r"^\d+(?:\.\d+){1,2}\s+", line):
            if kept:
                break
            continue
        if line in {"参考文献"}:
            break
        if line.startswith("图") or line.startswith(("–", "-", "—")):
            kept.append(line)
            continue
        if kept:
            break
    return "\n".join(kept).strip() if kept else text


def clean_cell_text(text: str) -> str:
    text = xml_safe_text(text)
    text = remove_ai_meta(text)
    text = strip_markdown_marks(text)
    text = enforce_terms(text)
    text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
    text = re.sub(r"\S+?\.qxd:[^\n]*?(?:Page\s*\d+|第\s*\d+\s*页)\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_paragraph(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return ""
    if all(re.match(r"^[•●○\-\u2022]", line) for line in lines):
        return "\n".join(lines)
    joined = "".join(lines)
    joined = re.sub(r"\s+", " ", joined)
    joined = re.sub(
        r"^(?:Long-Term Athlete Development Model|Physical Literacy|Sports Specialization|Trainability|Intellectual, Emotional, and Moral Development|Sport System Alignment and Integration|Stages of Long-Term Athlete Development)\s+3G E-LEARNING\s*",
        "",
        joined,
        flags=re.I,
    )
    joined = re.sub(r"^3G E-LEARNING\s+(?:Long-Term Athlete Development|长期运动员发展)\s*", "", joined, flags=re.I)
    joined = re.sub(r"\s+([，。；：？！、）])", r"\1", joined)
    joined = re.sub(r"([（])\s+", r"\1", joined)
    return joined.strip()


def is_standalone_subheading(text: str) -> bool:
    stripped = text.strip()
    if not (2 <= len(stripped) <= 34):
        return False
    if stripped.startswith(("*", "•", "●", "○", "-", "–")):
        return False
    if not re.search(r"[\u4e00-\u9fff]", stripped):
        return False
    if re.search(r"[。！？；，]$", stripped) or "；" in stripped:
        return False
    if re.search(r"\b(?:xv|xvi|xvii|xviii|xix|xx|xxi|xxii|xxiii)\b", stripped, flags=re.I):
        return False
    if re.search(r"\s+\d{1,3}$", stripped):
        return False
    if re.match(r"^(图|表)\s*\d", stripped):
        return False
    if stripped.count("·") >= 2:
        return False
    if re.search(r"(DOI|ISBN|国际标准书号|封面图片|标志设计|字体排版|排版公司|Routledge|Taylor\s*&\s*Francis)", stripped, flags=re.I):
        return False
    if re.search(r"(照片|摄影|Photo|Getty|Zuma|Icon\s+Sportswire|美联社|Associated\s+Press)", stripped, flags=re.I):
        return False
    if re.match(r"^\d+\s+[\u4e00-\u9fff·、与]+$", stripped):
        return False
    return True


def split_embedded_caption(text: str) -> tuple[str, str] | None:
    match = re.search(rf"({CHINESE_FIGURE_CAPTION_PATTERN})$", text.strip())
    if not match or match.start() == 0:
        return None
    before = text[: match.start()].strip()
    caption = match.group(1).strip()
    if not before or not caption:
        return None
    return before, caption


def is_figure_label_text(text: str) -> bool:
    stripped = re.sub(r"\s+", "", text.strip())
    if not stripped:
        return False
    if len(stripped) > 260:
        return False
    if re.search(r"[。！？；]", stripped):
        return False
    figure_terms = [
        "主教练", "董事会", "总经理", "技术总监", "竞技表现", "医疗部门", "营养部门",
        "物理治疗", "运动科学", "体能训练", "教练执照", "博士学位", "本科学位",
        "研究生学位", "专业发展", "认证", "欧足联", "亚足联", "美国足协",
    ]
    return sum(term in stripped for term in figure_terms) >= 2


def is_running_header_text(text: str) -> bool:
    stripped = re.sub(r"\s+", " ", text.strip())
    if not stripped or len(stripped) > 80:
        return False
    if "3G E-LEARNING" in stripped:
        return True
    if stripped in {"3G E-LEARNING", "长期运动员发展", "Long-Term Athlete Development"}:
        return True
    if re.match(r"^\d{1,3}\s+", stripped) and len(stripped) < 50 and not re.search(r"[,.，。:：]", stripped):
        return True
    if re.match(r"^\d{1,3}\s+[\u4e00-\u9fffA-Za-z·、\s]+(?:等|et al\.)?$", stripped, flags=re.I):
        return True
    if re.match(r"^[\u4e00-\u9fffA-Za-z ,&/\-]+?\s+\d{1,3}$", stripped):
        return True
    return False


def is_markdown_table(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return "|" in lines[index] and re.match(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$", lines[index + 1])


def is_loose_pipe_table(lines: list[str], index: int) -> bool:
    if index >= len(lines) or lines[index].count("|") < 1:
        return False
    seen_rows = 0
    i = index
    while i < len(lines) and lines[i].strip():
        if lines[i].count("|") >= 1:
            seen_rows += 1
        elif seen_rows:
            break
        i += 1
    return seen_rows >= 2


def parse_table(lines: list[str], index: int) -> tuple[TableBlock, int]:
    table_lines = []
    i = index
    while i < len(lines) and "|" in lines[i].strip():
        table_lines.append(lines[i].strip())
        i += 1
    rows: list[list[str]] = []
    for pos, line in enumerate(table_lines):
        if pos == 1 and re.match(r"^\s*\|?\s*:?-{2,}:?", line):
            continue
        cells = [clean_cell_text(cell.strip()) for cell in line.strip("|").split("|")]
        rows.append(cells)
    width = max((len(row) for row in rows), default=0)
    rows = [row + [""] * (width - len(row)) for row in rows if any(cell for cell in row)]
    return TableBlock(caption=None, rows=rows), i


def parse_loose_pipe_table(lines: list[str], index: int) -> tuple[TableBlock, int]:
    table_lines: list[str] = []
    i = index
    while i < len(lines) and lines[i].strip() and lines[i].count("|") >= 1:
        table_lines.append(lines[i].strip())
        i += 1
    rows: list[list[str]] = []
    for line in table_lines:
        if re.fullmatch(r"\|?\s*:?-{1,}:?\s*(\|\s*:?-{1,}:?\s*)+\|?", line):
            continue
        cells = [clean_cell_text(cell.strip()) for cell in line.strip("|").split("|")]
        if len(cells) >= 2 and any(cells):
            rows.append(cells)
    width = max((len(row) for row in rows), default=0)
    rows = [row + [""] * (width - len(row)) for row in rows if any(cell for cell in row)]
    return TableBlock(caption=None, rows=rows), i


def parse_inline_pipe_table(text: str) -> TableBlock | None:
    if text.count("|") < 4:
        return None
    raw_rows = [row.strip() for row in re.split(r"\s*\|\|\s*", text) if row.strip()]
    rows: list[list[str]] = []

    def is_separator_cell(cell: str) -> bool:
        return bool(re.fullmatch(r":?-{2,}:?", cell.replace(" ", "")))

    def is_numeric_cell(cell: str) -> bool:
        compact = cell.strip().replace(",", "")
        return bool(re.fullmatch(r"\d+(?:\.\d+)?(?:%|公斤|kg)?", compact, flags=re.IGNORECASE))

    def infer_width(cells: list[str]) -> int | None:
        for index, cell in enumerate(cells):
            if is_separator_cell(cell):
                return index if 2 <= index <= 12 else None
        for index, cell in enumerate(cells):
            if is_numeric_cell(cell):
                return index if 2 <= index <= 12 and len(cells) - index >= index else None
        for width in range(12, 1, -1):
            if len(cells) % width == 0 and len(cells) // width >= 2:
                return width
        return None

    if len(raw_rows) == 1:
        raw = raw_rows[0]
        cells = [clean_cell_text(cell.strip()) for cell in raw.strip("|").split("|")]
        cells = [cell for cell in cells if cell and not is_separator_cell(cell)]
        width = infer_width(cells)
        if width:
            rows = [cells[index : index + width] for index in range(0, len(cells), width)]
            rows = [row + [""] * (width - len(row)) for row in rows if any(row)]
    else:
        for raw in raw_rows:
            if re.fullmatch(r"\|?\s*:?-{1,}:?\s*(\|\s*:?-{1,}:?\s*)+\|?", raw):
                continue
            cells = [clean_cell_text(cell.strip()) for cell in raw.strip("|").split("|")]
            cells = [cell for cell in cells if not is_separator_cell(cell)]
            if len(cells) >= 2 and any(cells):
                rows.append(cells)
    if len(rows) < 2:
        return None
    width = max(len(row) for row in rows)
    if width < 2:
        return None
    rows = [row + [""] * (width - len(row)) for row in rows]
    return TableBlock(caption=None, rows=rows)


def classify_paragraph(text: str) -> str:
    stripped = text.strip()
    if (
        soccer_chapter_number(stripped) is not None
        or frank_chapter_number(stripped) is not None
        or gilbert_chapter_number(stripped) is not None
        or gordon_chapter_number(stripped) is not None
        or ltad_chapter_number(stripped) is not None
        or strength_chapter_number(stripped) is not None
    ):
        return "h1"
    if (
        re.match(r"^第\s*[一二三四五六七八九十\d]+章(?:[：:\s]+|$).{0,40}$", stripped)
        or re.match(r"^第\s*(?:[一二三四五六七八九十\d]+|[IVX]+)\s*部分(?:[：:\s]+|$).{0,40}$", stripped, flags=re.I)
        or re.match(r"^Chapter\s+\d+(?:\s+|$).{0,48}$", stripped, flags=re.I)
    ):
        return "h1"
    if re.match(r"^\d+\.\d+\.\d+\.?\s+[^\d\s].{0,70}$", stripped):
        return "h3"
    if re.match(r"^\d+\.\d+\.?\s+[^\d\s].{0,64}$", stripped):
        return "h2"
    if re.match(rf"^{CHINESE_FIGURE_CAPTION_PATTERN}$", stripped):
        return "caption"
    if re.match(r"^图\s*(?:[—–-]\s*.+|[（(]续[）)]\s*.*)$", stripped):
        return "caption"
    if re.match(r"^Figure\s+\d+(?:\.\d+)?[.．]?\s+.+", stripped, flags=re.I):
        return "caption"
    if re.match(rf"^{CHINESE_TABLE_CAPTION_PATTERN}$", stripped):
        return "table_caption"
    if re.match(r"^Table\s+\d+(?:\.\d+)?[.．]?\s*.*", stripped, flags=re.I):
        return "table_caption"
    if stripped in {"目录", "献辞", "前言", "序言", "引言", "参考文献", "Bibliografie", "Bibliography", "术语表", "作者简介", "关于作者", "致谢"}:
        return "h1"
    return "body"


def should_skip_source_toc_page(file: Path, text: str) -> bool:
    match = re.search(r"page_(\d+)", file.stem)
    page_number = int(match.group(1)) if match else None
    chapter_list_hits = len(re.findall(r"(?m)^\s*(?:第\s*[一二三四五六七八九十\d]+\s*章|第\s*[一二三四五六七八九十\dIVXivx]+\s*部分)\b", text))
    sentence_marks = len(re.findall(r"[。！？；]", text))
    short_lines = sum(1 for line in text.splitlines() if 1 <= len(line.strip()) <= 28)
    if chapter_list_hits >= 3 and short_lines >= chapter_list_hits and sentence_marks <= 2:
        return True
    if page_number and 3 <= page_number <= 15:
        compact = re.sub(r"\s+", "", text)
        toc_hits = sum(
            token in compact
            for token in ["目录", "表格目录", "图目录", "第一章", "第二章", "第三章", "第四章", "第五章", "第十章", "参考文献", "作者简介"]
        )
        number_lines = len(re.findall(r"(?m)^\s*\d{1,3}\s*$", text))
        dot_leaders = len(re.findall(r"[.…]{4,}", text))
        if page_number and 4 <= page_number <= 8 and ("目录" in compact or number_lines >= 8 or toc_hits >= 1):
            return True
        if page_number and 4 <= page_number <= 8 and short_lines >= 12 and sentence_marks <= 1:
            return True
        if page_number and 11 <= page_number <= 15 and ("目录" in compact or dot_leaders >= 3 or short_lines >= 18):
            return True
        if toc_hits >= 2 or dot_leaders >= 3 or (number_lines >= 8 and sentence_marks <= 2):
            return True
    return False


def should_skip_chapter_photo_text_page(file: Path, text: str) -> bool:
    match = re.search(r"page_(\d+)", file.stem)
    page_number = int(match.group(1)) if match else None
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    compact = re.sub(r"\s+", "", "\n".join(lines))
    if len(compact) > 120:
        return False
    if page_number and len(lines) <= 6 and any(re.fullmatch(r"第[一二三四五六七八九十\d]+章", line) for line in lines):
        if not any(re.match(r"^\d+\.\d+", line) for line in lines):
            return True
    return False


def is_chart_axis_noise(text: str) -> bool:
    stripped = text.strip()
    if len(stripped) > 140:
        return False
    if re.search(r"[，。；：？！]", stripped):
        return False
    decimal_count = len(re.findall(r"\d+\.\d+", stripped))
    month_count = len(re.findall(r"(?:十月|十二月|二月|四月|六月|January|February|March|April|May|June|July|August|September|October|November|December)", stripped, flags=re.I))
    if decimal_count >= 4 and month_count >= 2:
        return True
    if decimal_count >= 6 and len(re.findall(r"[\u4e00-\u9fff]", stripped)) <= 12:
        return True
    return False


def consume_chapter_heading(
    lines: list[str],
    index: int,
    chapter_title_overrides: dict[int, str] | None = None,
) -> tuple[str, int] | None:
    line = lines[index].strip()
    frank_single = frank_chapter_number(line) if chapter_title_overrides is FRANK_CHAPTER_TITLES else None
    if frank_single is not None:
        return FRANK_CHAPTER_TITLES.get(frank_single, line), index + 1
    numeric = re.match(r"^(\d{1,2})\s+(.+)$", line)
    if numeric:
        j = index + 1
        while j < len(lines) and not lines[j].strip():
            j += 1
        if j < len(lines):
            joined = f"{line}{lines[j].strip()}"
            frank_joined = frank_chapter_number(joined) if chapter_title_overrides is FRANK_CHAPTER_TITLES else None
            if frank_joined is not None:
                return FRANK_CHAPTER_TITLES.get(frank_joined, joined), j + 1
    if not re.fullmatch(r"第\s*[一二三四五六七八九十\d]+章", line):
        return None
    j = index + 1
    while j < len(lines) and not lines[j].strip():
        j += 1
    if j >= len(lines):
        return line, j
    title = lines[j].strip()
    if 2 <= len(title) <= 32 and not re.search(r"[。；：？！,，]|——", title):
        return f"{line} {title}", j + 1
    return line, index + 1


def bondarchuk_title_only(chapter: int) -> str:
    title = BONDARCHUK_CHAPTER_TITLES.get(chapter, "")
    return re.sub(r"^第[一二三四五六七八九十\d]+章\s*", "", title).strip()


def is_bondarchuk_generated_chapter_line(text: str) -> bool:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return False
    if re.match(r"^第\s*(?:X|x|[一二三四五六七八九十\d]+)\s*章(?:\s+|$)", stripped):
        return True
    return stripped in {bondarchuk_title_only(chapter) for chapter in BONDARCHUK_CHAPTER_TITLES}


def science_soccer_title_only(chapter: int) -> str:
    title = SCIENCE_SOCCER_CHAPTER_TITLES.get(chapter, "")
    return re.sub(r"^第[一二三四五六七八九十\d]+章\s*", "", title).strip()


def is_science_soccer_generated_heading_line(text: str) -> bool:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return False
    if re.match(r"^(?:SECTION|Section)\s+[A-F]\b", stripped):
        return True
    if re.match(r"^第\s*[A-FＡ-Ｆ]\s*部分\b", stripped):
        return True
    for title in SCIENCE_SOCCER_SECTION_TITLES.values():
        if stripped == title or stripped in title:
            return True
    for chapter, title in SCIENCE_SOCCER_CHAPTER_TITLES.items():
        title_only = science_soccer_title_only(chapter)
        if stripped == title or stripped == title_only:
            return True
        if re.match(rf"^{chapter}\s+", stripped) and (
            title_only in stripped
            or any(token in stripped.lower() for token in ["physical", "resistance", "aerobic", "soccer", "nutrition", "recovery", "psychological", "anticipation", "skill", "sociological", "wellbeing", "coaching", "injury", "infectious", "biomechanical", "analysis", "monitoring", "load", "growth", "talent", "scouting", "director"])
        ):
            return True
        if re.match(rf"^第\s*{chapter}\s*章\b", stripped):
            return True
    return False


def is_science_soccer_running_header(text: str) -> bool:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if re.fullmatch(r"(?:前言|序言)\s*[xivxlcdm\d]+", stripped, flags=re.I):
        return True
    if re.fullmatch(r"[xivxlcdm\d]+\s*(?:前言|序言)", stripped, flags=re.I):
        return True
    if re.fullmatch(r"\d{1,3}\s+[\u4e00-\u9fffA-Za-z·、\s]+", stripped) and len(stripped) <= 32:
        return True
    return False


def soviet_weightlifting_title_only(section: int) -> str:
    return SOVIET_WEIGHTLIFTING_TITLES.get(section, "").strip()


def is_soviet_weightlifting_generated_heading_line(text: str) -> bool:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return False
    title_set = {soviet_weightlifting_title_only(section) for section in SOVIET_WEIGHTLIFTING_TITLES}
    if stripped in title_set:
        return True
    aliases = {
        "本书适合哪些读者阅读",
        "本书适合哪些读者",
        "读者推荐",
        "读者推荐语",
        "致谢与前言",
        "什么是苏联体系？",
        "什么是苏联体系",
        "苏联体系的奠基人",
        "苏联体系的架构",
        "苏联体系的哲学基础",
        "提升竞技运动表现的策略与方法",
        "优化训练负荷的手段与方法",
        "制定训练计划",
        "十二周训练计划",
        "每周三次训练计划",
        "每周四次训练计划",
        "每周五次训练计划",
        "每周六次训练计划",
        "参考文献",
        "Works Cited",
        "作者简介",
        "关于作者",
        "关于里斯托体育",
        "关于 Risto Sports",
    }
    return stripped in aliases


def parse_translated_markdown(
    translated_dir: Path,
    images_by_page: dict[int, list[ImageBlock]] | None = None,
    source_reference_blocks: dict[int, list[Block]] | None = None,
    source_pages: dict[int, str] | None = None,
    use_standard_sections: bool = False,
    chapter_title_overrides: dict[int, str] | None = None,
) -> list[Block]:
    blocks: list[Block] = []
    pending_table_caption: str | None = None
    files = sorted(translated_dir.glob("page_*.md"))
    gordon_index_started = False
    for file in files:
        page_match = re.search(r"page_(\d+)", file.stem)
        page_number = int(page_match.group(1)) if page_match else None
        raw_text = filter_translation_artifact_chunks(file.read_text(encoding="utf-8", errors="ignore"))
        text = clean_text(raw_text)
        page_images = list(images_by_page.get(page_number, [])) if images_by_page and page_number else []
        source_refs_for_page = source_reference_blocks.get(page_number, []) if source_reference_blocks and page_number else []
        if chapter_title_overrides is SOVIET_WEIGHTLIFTING_TITLES and page_number in SOVIET_WEIGHTLIFTING_SKIP_PAGES:
            continue
        if chapter_title_overrides is PLATONOV_MOTOR_CHAPTER_TITLES and source_pages and page_number in source_pages:
            text = trim_low_text_figure_page(text, source_pages[page_number])
            text = trim_generated_figure_artifacts(text, source_pages[page_number])
        if source_refs_for_page:
            ref_match = re.search(r"(?im)^\s*(?:参考文献|References(?:\s+for\s+Chapter\s+\d+\.?)?)\s*$", text)
            if ref_match:
                text = text[: ref_match.start()].strip()
            else:
                blocks.extend(page_images)
                blocks.extend(source_refs_for_page)
                continue
        if chapter_title_overrides is PLATONOV_MOTOR_CHAPTER_TITLES and page_number in PLATONOV_MOTOR_SKIP_PAGES:
            continue
        if chapter_title_overrides is PLATONOV_MOTOR_CHAPTER_TITLES and page_number in PLATONOV_MOTOR_PART_START_PAGES:
            part = PLATONOV_MOTOR_PART_START_PAGES[page_number]
            blocks.append(ParagraphBlock(PLATONOV_MOTOR_PART_TITLES[part], "h1"))
            blocks.extend(page_images)
            continue
        if chapter_title_overrides is GILBERT_CHAPTER_TITLES and page_number in GILBERT_SKIP_PAGES:
            continue
        if chapter_title_overrides is GILBERT_CHAPTER_TITLES and source_pages and page_number in source_pages:
            text = trim_low_text_page_hallucination(text, source_pages[page_number])
            text = trim_generated_figure_artifacts(text, source_pages[page_number])
        if chapter_title_overrides is GILBERT_CHAPTER_TITLES and page_number in GILBERT_PART_TITLES:
            blocks.append(ParagraphBlock(GILBERT_PART_TITLES[page_number], "h1"))
            blocks.extend(page_images)
            continue
        if chapter_title_overrides is GORDON_CHAPTER_TITLES and page_number in GORDON_SKIP_PAGES:
            continue
        if chapter_title_overrides is GORDON_CHAPTER_TITLES and source_pages and page_number in source_pages:
            text = trim_low_text_page_hallucination(text, source_pages[page_number])
            text = trim_generated_figure_artifacts(text, source_pages[page_number])
        if chapter_title_overrides is GORDON_CHAPTER_TITLES and page_number in GORDON_PART_TITLES:
            blocks.append(ParagraphBlock(GORDON_PART_TITLES[page_number], "h1"))
            blocks.extend(page_images)
            continue
        if chapter_title_overrides is LTAD_CHAPTER_TITLES and page_number in LTAD_SKIP_PAGES:
            continue
        if chapter_title_overrides is LTAD_CHAPTER_TITLES and source_pages and page_number in source_pages:
            text = trim_low_text_page_hallucination(text, source_pages[page_number])
            text = trim_generated_figure_artifacts(text, source_pages[page_number])
        if chapter_title_overrides is LTAD_CHAPTER_TITLES and page_number and page_number < min(LTAD_CHAPTER_START_PAGES):
            text = re.sub(r"(?m)^(\s*)第\s*([一二三四五六七八九十\d]+)\s*章", r"\1本书第\2章", text)
        if chapter_title_overrides is LTAD_CHAPTER_TITLES and page_number in LTAD_FRONT_TITLES:
            blocks.append(ParagraphBlock(LTAD_FRONT_TITLES[page_number], "h1"))
        if chapter_title_overrides is GORDON_CHAPTER_TITLES and page_number and page_number >= 259:
            for raw_line in text.splitlines():
                stripped_line = normalize_paragraph(raw_line.strip())
                if not stripped_line:
                    continue
                if stripped_line in {"索引", "Index"}:
                    if not gordon_index_started:
                        blocks.append(ParagraphBlock("索引", "h1"))
                        gordon_index_started = True
                    continue
                if re.fullmatch(r"[A-Z]", stripped_line):
                    blocks.append(ParagraphBlock(stripped_line, "h2"))
                else:
                    blocks.append(ParagraphBlock(stripped_line, "index"))
            blocks.extend(page_images)
            continue
        if chapter_title_overrides is LTAD_CHAPTER_TITLES and page_number and page_number >= 325:
            for raw_line in text.splitlines():
                stripped_line = normalize_paragraph(raw_line.strip())
                if not stripped_line:
                    continue
                if is_running_header_text(stripped_line):
                    continue
                if stripped_line in {"索引", "Index"}:
                    blocks.append(ParagraphBlock("索引", "h1"))
                    continue
                if re.fullmatch(r"[A-Z]", stripped_line):
                    blocks.append(ParagraphBlock(stripped_line, "h2"))
                else:
                    blocks.append(ParagraphBlock(stripped_line, "index"))
            blocks.extend(page_images)
            continue
        if chapter_title_overrides is STRENGTH_CHAPTER_TITLES and page_number in STRENGTH_SKIP_PAGES:
            continue
        if chapter_title_overrides is STRENGTH_CHAPTER_TITLES and source_pages and page_number in source_pages:
            text = trim_low_text_page_hallucination(text, source_pages[page_number])
            text = trim_generated_figure_artifacts(text, source_pages[page_number])
        if chapter_title_overrides is STRENGTH_CHAPTER_TITLES and page_number and page_number < min(STRENGTH_CHAPTER_START_PAGES):
            text = re.sub(r"(?m)^(\s*)第\s*([一二三四五六七八九十\d]+)\s*章", r"\1本书第\2章", text)
        if chapter_title_overrides is STRENGTH_CHAPTER_TITLES and page_number in STRENGTH_FRONT_TITLES:
            blocks.append(ParagraphBlock(STRENGTH_FRONT_TITLES[page_number], "h1"))
        if chapter_title_overrides is STRENGTH_CHAPTER_TITLES and page_number in STRENGTH_PART_TITLES:
            blocks.append(ParagraphBlock(STRENGTH_PART_TITLES[page_number], "h1"))
            blocks.extend(page_images)
            continue
        if chapter_title_overrides is STRENGTH_CHAPTER_TITLES and page_number and 333 <= page_number <= 344:
            for raw_line in text.splitlines():
                stripped_line = normalize_paragraph(raw_line.strip())
                if not stripped_line:
                    continue
                if is_running_header_text(stripped_line):
                    continue
                if stripped_line in {"索引", "Index"}:
                    blocks.append(ParagraphBlock("索引", "h1"))
                    continue
                if re.fullmatch(r"[A-Z]", stripped_line):
                    blocks.append(ParagraphBlock(stripped_line, "h2"))
                else:
                    blocks.append(ParagraphBlock(stripped_line, "index"))
            blocks.extend(page_images)
            continue
        if chapter_title_overrides is BONDARCHUK_CHAPTER_TITLES and page_number in BONDARCHUK_SKIP_PAGES:
            continue
        if chapter_title_overrides is SCIENCE_SOCCER_CHAPTER_TITLES and page_number in SCIENCE_SOCCER_SKIP_PAGES:
            continue
        if should_skip_source_toc_page(file, text):
            continue
        if should_skip_chapter_photo_text_page(file, text):
            blocks.extend(page_images)
            continue
        lines = text.splitlines()
        paragraph_buffer: list[str] = []
        page_chapter_inserted = False
        if (
            chapter_title_overrides is SOVIET_WEIGHTLIFTING_TITLES
            and not lines
            and page_number in SOVIET_WEIGHTLIFTING_START_PAGES
        ):
            section = SOVIET_WEIGHTLIFTING_START_PAGES[page_number]
            blocks.append(ParagraphBlock(SOVIET_WEIGHTLIFTING_TITLES[section], "h1"))
            page_chapter_inserted = True
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if not line.strip():
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                i += 1
                continue
            if (
                chapter_title_overrides is SOVIET_WEIGHTLIFTING_TITLES
                and not page_chapter_inserted
                and page_number in SOVIET_WEIGHTLIFTING_START_PAGES
            ):
                section = SOVIET_WEIGHTLIFTING_START_PAGES[page_number]
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(SOVIET_WEIGHTLIFTING_TITLES[section], "h1"))
                page_chapter_inserted = True
                if is_soviet_weightlifting_generated_heading_line(line.strip()):
                    i += 1
                    continue
            if chapter_title_overrides is SOVIET_WEIGHTLIFTING_TITLES and is_soviet_weightlifting_generated_heading_line(line.strip()):
                i += 1
                continue
            if chapter_title_overrides is SCIENCE_SOCCER_CHAPTER_TITLES and is_science_soccer_running_header(line.strip()):
                i += 1
                continue
            if chapter_title_overrides is SCIENCE_SOCCER_CHAPTER_TITLES and page_number == 452:
                if not page_chapter_inserted:
                    if paragraph_buffer:
                        para = normalize_paragraph("\n".join(paragraph_buffer))
                        add_text_block(blocks, para, pending_table_caption, page_images)
                        if classify_paragraph(para) == "table_caption":
                            pending_table_caption = para
                        paragraph_buffer = []
                    blocks.append(ParagraphBlock("索引", "h1"))
                    page_chapter_inserted = True
                if re.search(r"\bIndex\b|索引", line.strip()) and len(line.strip()) <= 20:
                    i += 1
                    continue
            if chapter_title_overrides is SCIENCE_SOCCER_CHAPTER_TITLES and line.strip() in {"引言", "参考文献", "致谢", "结论", "总结", "未来方向与结论"}:
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(line.strip(), "h2"))
                i += 1
                continue
            if (
                chapter_title_overrides is SCIENCE_SOCCER_CHAPTER_TITLES
                and page_number in SCIENCE_SOCCER_SECTION_TITLES
                and not page_chapter_inserted
            ):
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(SCIENCE_SOCCER_SECTION_TITLES[page_number], "h1"))
                page_chapter_inserted = True
                if is_science_soccer_generated_heading_line(line.strip()):
                    i += 1
                    continue
            if (
                chapter_title_overrides is SCIENCE_SOCCER_CHAPTER_TITLES
                and not page_chapter_inserted
                and page_number in SCIENCE_SOCCER_CHAPTER_START_PAGES
            ):
                chapter = SCIENCE_SOCCER_CHAPTER_START_PAGES[page_number]
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(SCIENCE_SOCCER_CHAPTER_TITLES[chapter], "h1"))
                page_chapter_inserted = True
                if is_science_soccer_generated_heading_line(line.strip()):
                    i += 1
                    continue
            if chapter_title_overrides is SCIENCE_SOCCER_CHAPTER_TITLES and is_science_soccer_generated_heading_line(line.strip()):
                i += 1
                continue
            if (
                chapter_title_overrides is BONDARCHUK_CHAPTER_TITLES
                and not page_chapter_inserted
                and page_number in BONDARCHUK_CHAPTER_START_PAGES
            ):
                chapter = BONDARCHUK_CHAPTER_START_PAGES[page_number]
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(BONDARCHUK_CHAPTER_TITLES[chapter], "h1"))
                page_chapter_inserted = True
                if is_bondarchuk_generated_chapter_line(line.strip()):
                    i += 1
                    continue
            if chapter_title_overrides is BONDARCHUK_CHAPTER_TITLES and is_bondarchuk_generated_chapter_line(line.strip()):
                i += 1
                continue
            if (
                chapter_title_overrides is PLATONOV_MOTOR_CHAPTER_TITLES
                and page_chapter_inserted
                and page_number in PLATONOV_MOTOR_CHAPTER_START_PAGES
            ):
                chapter = PLATONOV_MOTOR_CHAPTER_START_PAGES[page_number]
                title_only = re.sub(r"^第[一二三四五六七八九十\d]+章\s*", "", PLATONOV_MOTOR_CHAPTER_TITLES[chapter]).strip()
                stripped_line = line.strip()
                if chinese_chapter_number(stripped_line) == chapter or stripped_line == title_only:
                    i += 1
                    continue
            chapter_heading = consume_chapter_heading(lines, i, chapter_title_overrides)
            if chapter_heading:
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                heading, i = chapter_heading
                blocks.append(ParagraphBlock(heading, "h1"))
                continue
            if (
                chapter_title_overrides is PLATONOV_MOTOR_CHAPTER_TITLES
                and not page_chapter_inserted
                and page_number in PLATONOV_MOTOR_CHAPTER_START_PAGES
            ):
                chapter = PLATONOV_MOTOR_CHAPTER_START_PAGES[page_number]
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(PLATONOV_MOTOR_CHAPTER_TITLES[chapter], "h1"))
                page_chapter_inserted = True
                title_only = re.sub(r"^第[一二三四五六七八九十\d]+章\s*", "", PLATONOV_MOTOR_CHAPTER_TITLES[chapter]).strip()
                stripped_line = line.strip()
                if chinese_chapter_number(stripped_line) == chapter or stripped_line == title_only or title_only in stripped_line:
                    i += 1
                    continue
            if (
                chapter_title_overrides is SOCCER_CHAPTER_TITLES
                and not page_chapter_inserted
                and page_number in SOCCER_CHAPTER_START_PAGES
                and SOCCER_CHAPTER_ALIASES.get(line.strip()) == SOCCER_CHAPTER_START_PAGES[page_number]
            ):
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(SOCCER_CHAPTER_TITLES[SOCCER_CHAPTER_START_PAGES[page_number]], "h1"))
                page_chapter_inserted = True
                i += 1
                continue
            if (
                chapter_title_overrides is GOMES_CHAPTER_TITLES
                and not page_chapter_inserted
                and page_number in GOMES_CHAPTER_START_PAGES
            ):
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(GOMES_CHAPTER_TITLES[GOMES_CHAPTER_START_PAGES[page_number]], "h1"))
                page_chapter_inserted = True
                if gomes_chapter_number(line.strip()) == GOMES_CHAPTER_START_PAGES[page_number]:
                    i += 1
                    continue
            if (
                chapter_title_overrides is GILBERT_CHAPTER_TITLES
                and not page_chapter_inserted
                and page_number in GILBERT_CHAPTER_START_PAGES
            ):
                chapter = GILBERT_CHAPTER_START_PAGES[page_number]
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(GILBERT_CHAPTER_TITLES[chapter], "h1"))
                page_chapter_inserted = True
                stripped_line = line.strip()
                if gilbert_chapter_number(stripped_line) == chapter:
                    i += 1
                    continue
                if re.fullmatch(r"第\s*[一二三四五六七八九十\d]+\s*章", stripped_line):
                    j = i + 1
                    while j < len(lines) and not lines[j].strip():
                        j += 1
                    if j < len(lines) and GILBERT_CHAPTER_ALIASES.get(lines[j].strip()) == chapter:
                        i = j + 1
                        continue
                if GILBERT_CHAPTER_ALIASES.get(stripped_line) == chapter:
                    i += 1
                    continue
            if (
                chapter_title_overrides is GORDON_CHAPTER_TITLES
                and not page_chapter_inserted
                and page_number in GORDON_CHAPTER_START_PAGES
            ):
                chapter = GORDON_CHAPTER_START_PAGES[page_number]
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(GORDON_CHAPTER_TITLES[chapter], "h1"))
                page_chapter_inserted = True
                stripped_line = line.strip()
                if gordon_chapter_number(stripped_line) == chapter:
                    i += 1
                    continue
                if re.fullmatch(r"第\s*[一二三四五六七八九十\d]+\s*章", stripped_line):
                    j = i + 1
                    while j < len(lines) and not lines[j].strip():
                        j += 1
                    if j < len(lines) and GORDON_CHAPTER_ALIASES.get(lines[j].strip()) == chapter:
                        i = j + 1
                        continue
                if GORDON_CHAPTER_ALIASES.get(stripped_line) == chapter:
                    i += 1
                    continue
            if (
                chapter_title_overrides is LTAD_CHAPTER_TITLES
                and not page_chapter_inserted
                and page_number in LTAD_CHAPTER_START_PAGES
            ):
                chapter = LTAD_CHAPTER_START_PAGES[page_number]
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(LTAD_CHAPTER_TITLES[chapter], "h1"))
                page_chapter_inserted = True
                stripped_line = line.strip()
                if ltad_chapter_number(stripped_line) == chapter:
                    i += 1
                    continue
                stripped_title = re.sub(r"^第\s*[一二三四五六七八九十\d]+\s*章[:：.\s]*", "", stripped_line).strip()
                if LTAD_CHAPTER_ALIASES.get(stripped_title) == chapter:
                    i += 1
                    continue
                if re.fullmatch(r"第\s*[一二三四五六七八九十\d]+\s*章", stripped_line):
                    j = i + 1
                    while j < len(lines) and not lines[j].strip():
                        j += 1
                    if j < len(lines) and LTAD_CHAPTER_ALIASES.get(lines[j].strip()) == chapter:
                        i = j + 1
                        continue
                if LTAD_CHAPTER_ALIASES.get(stripped_line) == chapter:
                    i += 1
                    continue
            if (
                chapter_title_overrides is STRENGTH_CHAPTER_TITLES
                and not page_chapter_inserted
                and page_number in STRENGTH_CHAPTER_START_PAGES
            ):
                chapter = STRENGTH_CHAPTER_START_PAGES[page_number]
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(STRENGTH_CHAPTER_TITLES[chapter], "h1"))
                page_chapter_inserted = True
                stripped_line = line.strip()
                if strength_chapter_number(stripped_line) == chapter:
                    i += 1
                    continue
                stripped_title = re.sub(r"^第\s*[一二三四五六七八九十\d]+\s*章[:：.\s]*", "", stripped_line).strip()
                if STRENGTH_CHAPTER_ALIASES.get(stripped_title) == chapter:
                    i += 1
                    continue
                if STRENGTH_CHAPTER_ALIASES.get(stripped_line) == chapter:
                    i += 1
                    continue
            if classify_paragraph(line.strip()) == "h1":
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                blocks.append(ParagraphBlock(line.strip(), "h1"))
                if chapter_title_overrides is SOCCER_CHAPTER_TITLES and page_number in SOCCER_CHAPTER_START_PAGES:
                    page_chapter_inserted = True
                if chapter_title_overrides is GOMES_CHAPTER_TITLES and page_number in GOMES_CHAPTER_START_PAGES:
                    page_chapter_inserted = True
                if chapter_title_overrides is GILBERT_CHAPTER_TITLES and page_number in GILBERT_CHAPTER_START_PAGES:
                    page_chapter_inserted = True
                if chapter_title_overrides is GORDON_CHAPTER_TITLES and page_number in GORDON_CHAPTER_START_PAGES:
                    page_chapter_inserted = True
                if chapter_title_overrides is LTAD_CHAPTER_TITLES and page_number in LTAD_CHAPTER_START_PAGES:
                    page_chapter_inserted = True
                if chapter_title_overrides is STRENGTH_CHAPTER_TITLES and page_number in STRENGTH_CHAPTER_START_PAGES:
                    page_chapter_inserted = True
                i += 1
                continue
            line_style = classify_paragraph(line.strip())
            if line_style in {"caption", "table_caption"}:
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    add_text_block(blocks, para, pending_table_caption, page_images)
                    if classify_paragraph(para) == "table_caption":
                        pending_table_caption = para
                    paragraph_buffer = []
                if line_style == "table_caption":
                    pending_table_caption = line.strip()
                else:
                    add_text_block(blocks, line.strip(), pending_table_caption, page_images)
                i += 1
                continue
            if is_markdown_table(lines, i):
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    style = classify_paragraph(para)
                    if style == "table_caption":
                        pending_table_caption = para
                    else:
                        blocks.append(ParagraphBlock(para, style))
                    paragraph_buffer = []
                table, i = parse_table(lines, i)
                table.caption = pending_table_caption
                pending_table_caption = None
                blocks.append(table)
                continue
            if is_loose_pipe_table(lines, i):
                group_caption = None
                if paragraph_buffer:
                    para = normalize_paragraph("\n".join(paragraph_buffer))
                    style = classify_paragraph(para)
                    if style == "table_caption":
                        pending_table_caption = para
                    elif len(para) <= 40 and not re.search(r"[。！？；]", para):
                        group_caption = para
                    else:
                        blocks.append(ParagraphBlock(para, style))
                    paragraph_buffer = []
                table, i = parse_loose_pipe_table(lines, i)
                if pending_table_caption and group_caption:
                    table.caption = f"{pending_table_caption}（{group_caption}）"
                else:
                    table.caption = group_caption or pending_table_caption
                pending_table_caption = None
                blocks.append(table)
                continue
            if is_standalone_subheading(line.strip()):
                group: list[str] = []
                j = i
                while j < len(lines) and lines[j].strip():
                    group.append(lines[j].strip())
                    j += 1
                if not any(classify_paragraph(item) in {"caption", "table_caption"} for item in group):
                    if paragraph_buffer:
                        para = normalize_paragraph("\n".join(paragraph_buffer))
                        add_text_block(blocks, para, pending_table_caption, page_images)
                        if classify_paragraph(para) == "table_caption":
                            pending_table_caption = para
                        paragraph_buffer = []
                    blocks.append(ParagraphBlock(line.strip(), "h2"))
                    i += 1
                    continue
            paragraph_buffer.append(line)
            i += 1
        if paragraph_buffer:
            para = normalize_paragraph("\n".join(paragraph_buffer))
            add_text_block(blocks, para, pending_table_caption, page_images)
            if classify_paragraph(para) == "table_caption":
                pending_table_caption = para
        if page_images:
            blocks.extend(page_images)
        if source_refs_for_page:
            blocks.extend(source_refs_for_page)
    parsed_blocks = sanitize_heading_blocks(
        merge_small_body_blocks(blocks),
        use_standard_sections=use_standard_sections,
        chapter_title_overrides=chapter_title_overrides,
    )
    if chapter_title_overrides is GILBERT_CHAPTER_TITLES:
        parsed_blocks = fix_gilbert_specific_blocks(parsed_blocks)
    return parsed_blocks


def chinese_chapter_number(text: str) -> int | None:
    match = re.match(r"^第\s*([一二三四五六七八九十\dIVXivx]+)\s*章", text.strip())
    if not match:
        return None
    raw = match.group(1)
    return chinese_number_value(raw)


def chinese_number_value(raw: str) -> int | None:
    if raw.isdigit():
        return int(raw)
    roman_values = {"I": 1, "V": 5, "X": 10}
    roman = raw.upper()
    if roman and all(char in roman_values for char in roman):
        total = 0
        prev = 0
        for char in reversed(roman):
            value = roman_values[char]
            if value < prev:
                total -= value
            else:
                total += value
                prev = value
        return total or None
    values = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
    if raw == "十":
        return 10
    if raw.startswith("十"):
        return 10 + values.get(raw[1:], 0)
    if raw.endswith("十"):
        return values.get(raw[0], 0) * 10
    if "十" in raw:
        left, right = raw.split("十", 1)
        return values.get(left, 1) * 10 + values.get(right, 0)
    return values.get(raw)


def chinese_part_number(text: str) -> int | None:
    match = re.match(r"^第?\s*([一二三四五六七八九十\dIVXivx]+)\s*部分\b", text.strip())
    if not match:
        return None
    return chinese_number_value(match.group(1))


def soccer_chapter_number(text: str) -> int | None:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return None
    for title, number in SOCCER_CHAPTER_ALIASES.items():
        if stripped == title and stripped not in SOCCER_GENERIC_BARE_HEADINGS:
            return number
    numeric = re.match(r"^(\d{1,2})\s+(.+)$", stripped)
    if numeric:
        number = int(numeric.group(1))
        title = numeric.group(2).strip()
        if SOCCER_CHAPTER_ALIASES.get(title) == number:
            return number
    return None


def frank_chapter_number(text: str) -> int | None:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return None
    match = re.match(r"^第\s*(\d{1,2})\s*章\s*(.+)?$", stripped)
    if match:
        number = int(match.group(1))
        title = (match.group(2) or "").strip()
        if number in FRANK_CHAPTER_TITLES and (not title or FRANK_CHAPTER_ALIASES.get(title) == number or title in FRANK_CHAPTER_TITLES[number]):
            return number
    match = re.match(r"^(\d{1,2})\s+(.+)$", stripped)
    if match:
        number = int(match.group(1))
        title = match.group(2).strip()
        if FRANK_CHAPTER_ALIASES.get(title) == number:
            return number
    return None


def gomes_chapter_number(text: str) -> int | None:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return None
    match = re.match(r"^第\s*(\d{1,2})\s*章\s*(.+)$", stripped)
    if match:
        number = int(match.group(1))
        title = match.group(2).strip(" ：:")
        if GOMES_CHAPTER_ALIASES.get(title) == number:
            return number
        return None
    match = re.match(r"^(\d{1,2})[.、]?\s+(.+)$", stripped)
    if match:
        number = int(match.group(1))
        title = match.group(2).strip(" ：:")
        if GOMES_CHAPTER_ALIASES.get(title) == number:
            return number
    return GOMES_CHAPTER_ALIASES.get(stripped)


def gilbert_chapter_number(text: str) -> int | None:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return None
    number = chinese_chapter_number(stripped)
    if number in GILBERT_CHAPTER_TITLES:
        title = re.sub(r"^第\s*[一二三四五六七八九十\d]+\s*章[:：.\s]*", "", stripped).strip()
        if not title or GILBERT_CHAPTER_ALIASES.get(title) == number:
            return number
    match = re.match(r"^第\s*(\d{1,2})\s*章[:：.\s]*(.+)?$", stripped)
    if match:
        number = int(match.group(1))
        title = (match.group(2) or "").strip()
        if number in GILBERT_CHAPTER_TITLES and (not title or GILBERT_CHAPTER_ALIASES.get(title) == number):
            return number
    match = re.match(r"^(\d{1,2})[.、]?\s+(.+)$", stripped)
    if match:
        number = int(match.group(1))
        title = match.group(2).strip(" ：:")
        if GILBERT_CHAPTER_ALIASES.get(title) == number:
            return number
    return GILBERT_CHAPTER_ALIASES.get(stripped)


def gordon_chapter_number(text: str) -> int | None:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return None
    number = chinese_chapter_number(stripped)
    if number in GORDON_CHAPTER_TITLES:
        title = re.sub(r"^第\s*[一二三四五六七八九十\d]+\s*章[:：.\s]*", "", stripped).strip()
        if not title or GORDON_CHAPTER_ALIASES.get(title) == number:
            return number
    match = re.match(r"^第\s*(\d{1,2})\s*章[:：.\s]*(.+)?$", stripped)
    if match:
        number = int(match.group(1))
        title = (match.group(2) or "").strip()
        if number in GORDON_CHAPTER_TITLES and (not title or GORDON_CHAPTER_ALIASES.get(title) == number):
            return number
    match = re.match(r"^(\d{1,2})[.、]?\s+(.+)$", stripped)
    if match:
        number = int(match.group(1))
        title = match.group(2).strip(" ：:")
        if GORDON_CHAPTER_ALIASES.get(title) == number:
            return number
    return GORDON_CHAPTER_ALIASES.get(stripped)


def ltad_chapter_number(text: str) -> int | None:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return None
    number = chinese_chapter_number(stripped)
    if number in LTAD_CHAPTER_TITLES:
        title = re.sub(r"^第\s*[一二三四五六七八九十\d]+\s*章[:：.\s]*", "", stripped).strip()
        alias_number = LTAD_CHAPTER_ALIASES.get(title)
        if alias_number is not None:
            return alias_number
        if not title:
            return number
    match = re.match(r"^第\s*(\d{1,2})\s*章[:：.\s]*(.+)?$", stripped)
    if match:
        number = int(match.group(1))
        title = (match.group(2) or "").strip()
        alias_number = LTAD_CHAPTER_ALIASES.get(title)
        if alias_number is not None:
            return alias_number
        if number in LTAD_CHAPTER_TITLES and not title:
            return number
    match = re.match(r"^(\d{1,2})[.、]?\s+(.+)$", stripped)
    if match:
        number = int(match.group(1))
        title = match.group(2).strip(" ：:")
        if LTAD_CHAPTER_ALIASES.get(title) == number:
            return number
    return None


def strength_chapter_number(text: str) -> int | None:
    stripped = re.sub(r"[*#`]+", "", text).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if not stripped:
        return None
    number = chinese_chapter_number(stripped)
    if number in STRENGTH_CHAPTER_TITLES:
        title = re.sub(r"^第\s*[一二三四五六七八九十\d]+\s*章[:：.\s]*", "", stripped).strip()
        alias_number = STRENGTH_CHAPTER_ALIASES.get(title)
        if alias_number is not None:
            return alias_number
        if not title:
            return number
    match = re.match(r"^第\s*(\d{1,2})\s*章[:：.\s]*(.+)?$", stripped)
    if match:
        number = int(match.group(1))
        title = (match.group(2) or "").strip()
        alias_number = STRENGTH_CHAPTER_ALIASES.get(title)
        if alias_number is not None:
            return alias_number
        if number in STRENGTH_CHAPTER_TITLES and not title:
            return number
    match = re.match(r"^(\d{1,2})[.、]?\s+(.+)$", stripped)
    if match:
        number = int(match.group(1))
        title = match.group(2).strip(" ：:")
        if STRENGTH_CHAPTER_ALIASES.get(title) == number:
            return number
    return None


def heading_number(text: str) -> tuple[int, ...] | None:
    match = re.match(r"^(\d+)(?:\.(\d+))?(?:\.(\d+))?", text.strip())
    if not match:
        return None
    return tuple(int(part) for part in match.groups() if part is not None)


def heading_title_text(text: str) -> str:
    return re.sub(r"^\d+(?:\.\d+){0,2}\.?\s*", "", text.strip())


def title_similarity(candidate: str, expected: str) -> float:
    candidate = re.sub(r"[^\u4e00-\u9fffA-Za-z0-9]", "", candidate)
    expected = re.sub(r"[^\u4e00-\u9fffA-Za-z0-9]", "", expected)
    if not candidate or not expected:
        return 0.0
    if candidate in expected or expected in candidate:
        return 1.0
    expected_pairs = {expected[i : i + 2] for i in range(max(len(expected) - 1, 0))}
    candidate_pairs = {candidate[i : i + 2] for i in range(max(len(candidate) - 1, 0))}
    if not expected_pairs:
        return 0.0
    return len(expected_pairs & candidate_pairs) / len(expected_pairs)


def sanitize_heading_blocks(
    blocks: list[Block],
    use_standard_sections: bool = False,
    chapter_title_overrides: dict[int, str] | None = None,
) -> list[Block]:
    sanitized: list[Block] = []
    current_chapter: int | None = None
    seen_numbers: dict[tuple[int, ...], int] = {}
    seen_scores: dict[tuple[int, ...], float] = {}
    seen_h1_titles: set[str] = set()
    seen_h1_chapters: set[int] = set()
    seen_h1_parts: set[int] = set()
    last_h1_title: str | None = None
    for block in blocks:
        if not isinstance(block, ParagraphBlock):
            sanitized.append(block)
            continue
        text = block.text.strip()
        if is_running_header_text(text):
            continue
        if block.style in {"body", "h2"} and last_h1_title and text == last_h1_title:
            continue
        if block.style == "h1":
            part = chinese_part_number(text)
            if part is not None:
                if part > 8 or re.search(r"第\s*X\s*部分", text, flags=re.I):
                    block.style = "body"
                elif part in seen_h1_parts:
                    block.style = "body"
                else:
                    seen_h1_parts.add(part)
                    if chapter_title_overrides is PLATONOV_CHAPTER_TITLES:
                        block.text = PLATONOV_PART_TITLES.get(part, block.text)
                        text = block.text.strip()
            if chapter_title_overrides is PLATONOV_MOTOR_CHAPTER_TITLES and part is not None:
                block.text = PLATONOV_MOTOR_PART_TITLES.get(part, block.text)
                text = block.text.strip()
            if chapter_title_overrides is LTAD_CHAPTER_TITLES and re.fullmatch(
                r"第\s*[一二三四五六七八九十\d]+\s*章[:：.\s]*引言", text
            ):
                block.text = "引言"
                block.style = "h2"
                sanitized.append(block)
                continue
            if chapter_title_overrides is GOMES_CHAPTER_TITLES:
                chapter = gomes_chapter_number(text)
                if chapter is None and text not in {"致谢", "序言", "前言", "引言", "参考文献", "结语"}:
                    block.style = "body"
                    sanitized.append(block)
                    continue
            else:
                chapter = chinese_chapter_number(text)
                if chapter is None and chapter_title_overrides:
                    chapter = soccer_chapter_number(text)
                if chapter is None and chapter_title_overrides is FRANK_CHAPTER_TITLES:
                    chapter = frank_chapter_number(text)
                if chapter is None and chapter_title_overrides is GILBERT_CHAPTER_TITLES:
                    chapter = gilbert_chapter_number(text)
                if chapter is None and chapter_title_overrides is GORDON_CHAPTER_TITLES:
                    chapter = gordon_chapter_number(text)
                if chapter is None and chapter_title_overrides is LTAD_CHAPTER_TITLES:
                    chapter = ltad_chapter_number(text)
                if chapter is None and chapter_title_overrides is STRENGTH_CHAPTER_TITLES:
                    chapter = strength_chapter_number(text)
            if chapter is not None:
                if re.search(r"\d+\.\d+", text):
                    block.style = "body"
                elif chapter in seen_h1_chapters:
                    block.style = "body"
                else:
                    seen_h1_chapters.add(chapter)
                    if chapter_title_overrides:
                        block.text = chapter_title_overrides.get(chapter, block.text)
                    text = block.text.strip()
                    current_chapter = chapter
                    if text in seen_h1_titles:
                        block.style = "body"
            elif text in {"参考文献"} and current_chapter is not None:
                block.style = "h2"
            elif chapter_title_overrides is GILBERT_CHAPTER_TITLES and text in {
                "献辞",
                "致谢",
                "引言",
                "参考文献",
                "作者简介",
                "关于作者",
                "结语：重复成功",
                *GILBERT_PART_TITLES.values(),
            }:
                if text in seen_h1_titles:
                    block.style = "body"
            elif chapter_title_overrides is GORDON_CHAPTER_TITLES and text in {
                "献辞",
                "前言",
                "致谢",
                "参考文献",
                "索引",
                *GORDON_PART_TITLES.values(),
            }:
                if text in seen_h1_titles:
                    block.style = "body"
            elif chapter_title_overrides is LTAD_CHAPTER_TITLES and text in {
                "版权信息",
                "编辑委员会",
                "前言",
                "参考文献",
                "索引",
            }:
                if text in seen_h1_titles:
                    block.style = "body"
            elif chapter_title_overrides is LTAD_CHAPTER_TITLES:
                if text == "引言":
                    block.style = "h2"
                else:
                    block.style = "body"
            elif chapter_title_overrides is STRENGTH_CHAPTER_TITLES and text in {
                "前言",
                "序言",
                "致谢",
                "符号与缩略语",
                "术语表",
                "参考书目",
                "Bibliography",
                "索引",
                "关于作者",
                *STRENGTH_PART_TITLES.values(),
            }:
                if text in seen_h1_titles:
                    block.style = "body"
            elif chapter_title_overrides is STRENGTH_CHAPTER_TITLES:
                block.style = "body"
            elif chapter_title_overrides is PLATONOV_MOTOR_CHAPTER_TITLES and text in {
                "缩略语表",
                "作者的话",
                "作者序",
                "参考文献",
                *PLATONOV_MOTOR_PART_TITLES.values(),
            }:
                if text in seen_h1_titles:
                    block.style = "body"
            elif chapter_title_overrides is PLATONOV_MOTOR_CHAPTER_TITLES:
                block.style = "body"
            elif text in seen_h1_titles:
                block.style = "body"
            if block.style == "h1":
                seen_h1_titles.add(text)
                last_h1_title = text
            sanitized.append(block)
            continue
        if block.style in {"h2", "h3"}:
            number = heading_number(text)
            if chapter_title_overrides is STRENGTH_CHAPTER_TITLES and current_chapter is None and (number or text.startswith("本书第")):
                block.style = "body"
            if current_chapter and number and number[0] != current_chapter:
                block.style = "body"
            if use_standard_sections and number and number not in ALLOWED_SECTION_NUMBERS:
                block.style = "body"
            if re.search(r"(十月|十二月|二月|四月|六月)", text):
                block.style = "body"
            if block.style in {"h2", "h3"} and number:
                expected_title = STANDARD_SECTION_TITLES.get(number) if use_standard_sections else None
                score = title_similarity(heading_title_text(text), expected_title or "")
                if expected_title and score < 0.25:
                    block.style = "body"
                else:
                    block.text = f"{'.'.join(str(part) for part in number)} {expected_title}" if expected_title else block.text
                previous_index = seen_numbers.get(number)
                if block.style in {"h2", "h3"} and previous_index is not None and isinstance(sanitized[previous_index], ParagraphBlock):
                    previous = sanitized[previous_index]
                    previous_score = seen_scores.get(number, 0.0)
                    if score > previous_score + 0.05:
                        previous.style = "body"
                        seen_numbers[number] = len(sanitized)
                        seen_scores[number] = score
                    else:
                        block.style = "body"
                elif block.style in {"h2", "h3"}:
                    seen_numbers[number] = len(sanitized)
                    seen_scores[number] = score
        sanitized.append(block)
    cleaned: list[Block] = []
    recent_h1: str | None = None
    for block in sanitized:
        if isinstance(block, ParagraphBlock):
            block_text = block.text.strip()
            if block.style == "h1":
                recent_h1 = block_text
            elif block.style in {"body", "h2"} and (block_text in seen_h1_titles or (recent_h1 and block_text == recent_h1)):
                continue
        cleaned.append(block)
    return cleaned


def fix_gilbert_specific_blocks(blocks: list[Block]) -> list[Block]:
    fixed: list[Block] = []
    index = 0
    while index < len(blocks):
        block = blocks[index]
        if (
            isinstance(block, ParagraphBlock)
            and block.style == "body"
            and block.text.rstrip().endswith("这段陈述是")
        ):
            block.text = block.text.rstrip()[: -len("这段陈述是")] + "这段陈述并不是详细阐述执教哲学的地方。"
            fixed.append(block)
            index += 1
            if (
                index < len(blocks)
                and isinstance(blocks[index], ParagraphBlock)
                and blocks[index].style in {"h2", "h3"}
                and blocks[index].text.strip() == "执教目的与核心价值观"
            ):
                index += 1
            if (
                index < len(blocks)
                and isinstance(blocks[index], ParagraphBlock)
                and blocks[index].style == "body"
                and blocks[index].text.startswith("这里并非阐述你执教哲学的场所。")
            ):
                remainder = blocks[index].text.replace("这里并非阐述你执教哲学的场所。", "", 1).strip()
                if remainder:
                    fixed.append(ParagraphBlock(remainder, "body"))
                index += 1
            continue
        fixed.append(block)
        index += 1
    return fixed


def add_text_block(
    blocks: list[Block],
    para: str,
    pending_table_caption: str | None,
    page_images: list[ImageBlock] | None = None,
) -> None:
    if not para:
        return
    if is_chart_axis_noise(para):
        return
    if is_running_header_text(para):
        return
    caption_split = split_embedded_caption(para)
    if caption_split:
        before, caption = caption_split
        if is_figure_label_text(before):
            if page_images:
                image = page_images.pop(0)
                image.caption = caption
                blocks.append(image)
            else:
                blocks.append(ParagraphBlock(caption, "caption"))
            return
        add_text_block(blocks, before, pending_table_caption, page_images)
        add_text_block(blocks, caption, pending_table_caption, page_images)
        return
    inline_table = parse_inline_pipe_table(para)
    if inline_table is not None:
        inline_table.caption = pending_table_caption
        blocks.append(inline_table)
        return
    style = classify_paragraph(para)
    if style == "table_caption":
        return
    if style == "caption" and page_images:
        image = page_images.pop(0)
        image.caption = para
        blocks.append(image)
        return
    blocks.append(ParagraphBlock(para, style))


def merge_small_body_blocks(blocks: list[Block]) -> list[Block]:
    merged: list[Block] = []
    for block in blocks:
        if isinstance(block, ParagraphBlock):
            inline_table = parse_inline_pipe_table(block.text)
            if inline_table is not None:
                merged.append(inline_table)
                continue
        if (
            isinstance(block, ParagraphBlock)
            and block.style == "body"
            and merged
            and isinstance(merged[-1], ParagraphBlock)
            and merged[-1].style == "body"
            and len(merged[-1].text) < 80
            and not re.search(r"[。？！：；]$", merged[-1].text)
        ):
            merged[-1].text = merged[-1].text + block.text
        else:
            merged.append(block)
    return merged


def configure_doc(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.footer_distance = Cm(1.4)

    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    styles = {
        "Heading 1": (HEADING_FONT, Pt(16)),
        "Heading 2": (HEADING_FONT, Pt(14)),
        "Heading 3": (HEADING_FONT, Pt(12)),
    }
    for name, (font, size) in styles.items():
        style = document.styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
    for toc_name in ["TOC 1", "TOC 2", "TOC 3"]:
        if toc_name in document.styles:
            style = document.styles[toc_name]
            style.font.name = LATIN_FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
            style.font.color.rgb = RGBColor(0, 0, 0)


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])
    set_run_font(run, BODY_FONT, LATIN_FONT, Pt(9))


def add_cover(document: Document, title: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    run = p.add_run(title)
    set_run_font(run, HEADING_FONT, LATIN_FONT, Pt(24))
    run.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("中文译制版")
    set_run_font(run, BODY_FONT, LATIN_FONT, Pt(14))
    document.add_section(WD_SECTION.NEW_PAGE)


def add_translation_note(document: Document) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(36)
    heading.paragraph_format.space_after = Pt(18)
    run = heading.add_run("译制说明")
    set_run_font(run, HEADING_FONT, LATIN_FONT, Pt(16))
    run.bold = True

    note = (
        "本译制稿依据原书内容整理，保留原作者、出版信息、参考文献及索引等必要书目信息。"
        "正文术语按运动训练学、教练科学与竞技体育训练领域的中文表达统一处理，"
        "图表尽量按原书关系回插，以便学习研究和内部阅读使用。"
    )
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(note)
    set_run_font(run, BODY_FONT, LATIN_FONT, Pt(12))
    document.add_section(WD_SECTION.NEW_PAGE)


def add_toc(document: Document) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(12)
    run = heading.add_run("目录")
    set_run_font(run, HEADING_FONT, LATIN_FONT, Pt(16))
    run.bold = True
    p = document.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-1" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])
    document.add_section(WD_SECTION.NEW_PAGE)


def add_paragraph(document: Document, block: ParagraphBlock) -> None:
    if block.style == "h1":
        text = paragraph_display_text(block.text)
        p = document.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)
        for run in p.runs:
            set_run_font(run, HEADING_FONT, LATIN_FONT, Pt(16))
            run.font.color.rgb = RGBColor(0, 0, 0)
        return
    if block.style == "h2":
        text = paragraph_display_text(block.text)
        p = document.add_heading(text, level=2)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        set_keep_with_next(p)
        for run in p.runs:
            set_run_font(run, HEADING_FONT, LATIN_FONT, Pt(14))
            run.font.color.rgb = RGBColor(0, 0, 0)
        return
    if block.style == "h3":
        text = paragraph_display_text(block.text)
        p = document.add_heading(text, level=3)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        set_keep_with_next(p)
        for run in p.runs:
            set_run_font(run, HEADING_FONT, LATIN_FONT, Pt(12))
            run.font.color.rgb = RGBColor(0, 0, 0)
        return
    if block.style == "caption":
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(paragraph_display_text(block.text))
        set_run_font(run, BODY_FONT, LATIN_FONT, Pt(10.5))
        return
    if block.style == "reference":
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(paragraph_display_text(block.text))
        set_run_font(run, LATIN_FONT, LATIN_FONT, Pt(11))
        return
    if block.style == "index":
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(paragraph_display_text(block.text))
        set_run_font(run, BODY_FONT, LATIN_FONT, Pt(10.5))
        return
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(paragraph_display_text(block.text))
    set_run_font(run, BODY_FONT, LATIN_FONT, Pt(12))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(document: Document, block: TableBlock) -> None:
    if block.caption:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(paragraph_display_text(block.caption))
        set_run_font(run, BODY_FONT, LATIN_FONT, Pt(10.5))
    if not block.rows:
        return
    rows, cols = len(block.rows), max(len(row) for row in block.rows)
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for r, row in enumerate(block.rows):
        for c in range(cols):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = paragraph_display_text(row[c] if c < len(row) else "")
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, BODY_FONT, LATIN_FONT, Pt(9.5))
            if r == 0:
                run.bold = True
    document.add_paragraph()


def add_image(document: Document, block: ImageBlock) -> None:
    path = Path(block.path)
    if not path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(path), width=Cm(13.5))
    except Exception:
        try:
            run.add_picture(str(path), width=Inches(5.0))
        except Exception:
            return
    if not block.caption:
        return
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run(paragraph_display_text(block.caption))
    set_run_font(run, BODY_FONT, LATIN_FONT, Pt(10.5))


def render_docx(blocks: list[Block], output_docx: Path, title: str) -> None:
    document = Document()
    configure_doc(document)
    add_page_number(document.sections[0])
    add_cover(document, title)
    add_translation_note(document)
    add_toc(document)
    for block in blocks:
        if isinstance(block, TableBlock):
            add_table(document, block)
        elif isinstance(block, ImageBlock):
            add_image(document, block)
        elif isinstance(block, ParagraphBlock):
            add_paragraph(document, block)
    remove_trailing_empty_paragraphs(document)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def remove_trailing_empty_paragraphs(document: Document) -> None:
    for paragraph in reversed(document.paragraphs):
        if paragraph.text.strip():
            break
        paragraph._element.getparent().remove(paragraph._element)


def trim_trailing_blank_pdf_pages(pdf_path: Path) -> None:
    with fitz.open(pdf_path) as doc:
        delete_from = doc.page_count
        while delete_from > 1:
            page = doc[delete_from - 1]
            text = re.sub(r"\s+", "", page.get_text("text"))
            if text and not re.fullmatch(r"\d+", text):
                break
            if page.get_images(full=True):
                break
            delete_from -= 1
        if delete_from == doc.page_count:
            return
        for index in range(doc.page_count - 1, delete_from - 1, -1):
            doc.delete_page(index)
        tmp_pdf = pdf_path.with_suffix(".trimmed.pdf")
        doc.save(tmp_pdf)
    tmp_pdf.replace(pdf_path)


def export_pdf_with_word(docx_path: Path, pdf_path: Path) -> None:
    script = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{str(docx_path).replace("'", "''")}')
$doc.Fields.Update() | Out-Null
foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() | Out-Null }}
$doc.Save()
$doc.SaveAs([ref] '{str(pdf_path).replace("'", "''")}', [ref] 17)
$doc.Close()
$word.Quit()
"""
    subprocess.run(["powershell", "-NoProfile", "-Command", script], check=True)
    trim_trailing_blank_pdf_pages(pdf_path)


def validate_final_output(output_dir: Path) -> list[str]:
    errors: list[str] = []
    docx = output_dir / "final_zh.docx"
    pdf = output_dir / "final_zh.pdf"
    if not docx.exists():
        errors.append("final_zh.docx missing")
    if not pdf.exists():
        errors.append("final_zh.pdf missing")
    extra = [p.name for p in output_dir.iterdir() if p.name not in {"final_zh.docx", "final_zh.pdf"}]
    if extra:
        errors.append("extra files in final output: " + ", ".join(extra))
    if docx.exists():
        # DOCX is zipped XML; these checks catch accidental draft markers in document text.
        import zipfile

        with zipfile.ZipFile(docx) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        plain_xml = html.unescape(xml)
        for token in ["—— 原书第", "<br>", "原书第 X 页", "原书第", "第 X 章", "第X章", "第 X 部分", "第X部分", "认知性竞技战术", "认知性专项战术", "认知性技术技能", "超量补偿"]:
            if token in plain_xml:
                errors.append(f"forbidden marker found: {token}")
        if re.search(r"\|[^|\n]+\|", plain_xml):
            errors.append("markdown table marker found")
        if re.search(r"3\.5\s*3\.6\s*3\.7.*十月.*十二月", plain_xml):
            errors.append("chart axis text leaked into document body")
        if len(Document(docx).inline_shapes) == 0:
            errors.append("no images were inserted")
    return errors


def clean_output_dir(output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    for child in output_dir.iterdir():
        if child.is_dir():
            shutil.rmtree(child)
        else:
            child.unlink()


def build_image_blocks(images: list[ExtractedImage]) -> dict[int, list[ImageBlock]]:
    images_by_page: dict[int, list[ImageBlock]] = {}
    figure_index = 1
    for image in images:
        if image.width < 120 or image.height < 120:
            continue
        ratio = max(image.width / max(image.height, 1), image.height / max(image.width, 1))
        if ratio > 8:
            continue
        image_path = normalize_image_for_docx(Path(image.path))
        if image_path is None:
            continue
        caption = f"图 {figure_index} 书中插图"
        images_by_page.setdefault(image.page_number, []).append(ImageBlock(image_path, caption))
        figure_index += 1
    return images_by_page


def union_rects(rects: list[fitz.Rect]) -> fitz.Rect | None:
    if not rects:
        return None
    union = fitz.Rect(rects[0])
    for rect in rects[1:]:
        union |= rect
    return union


def expand_rect(rect: fitz.Rect, page_rect: fitz.Rect, margin: float = 10) -> fitz.Rect:
    return fitz.Rect(
        max(page_rect.x0, rect.x0 - margin),
        max(page_rect.y0, rect.y0 - margin),
        min(page_rect.x1, rect.x1 + margin),
        min(page_rect.y1, rect.y1 + margin),
    )


def render_clip(page, clip: fitz.Rect, image_path: Path) -> None:
    pix = page.get_pixmap(matrix=fitz.Matrix(3.0, 3.0), clip=clip, alpha=False)
    pix.save(image_path)


def caption_rects(page) -> list[tuple[str, fitz.Rect]]:
    found: list[tuple[str, fitz.Rect]] = []
    for block in page.get_text("blocks"):
        x0, y0, x1, y1, text, *_ = block
        match = SOURCE_FIGURE_CAPTION_RE.match(text or "")
        if match:
            found.append((match.group(1), fitz.Rect(x0, y0, x1, y1)))
    found.sort(key=lambda item: (item[1].y0, item[1].x0))
    return found


def figure_clip_for_caption(page, caption_rect: fitz.Rect) -> fitz.Rect | None:
    limit_y = caption_rect.y0 - 4
    lower_y = max(0, caption_rect.y0 - 420)
    rects: list[fitz.Rect] = []
    for drawing in page.get_drawings():
        rect = drawing.get("rect")
        if rect and rect.width > 8 and rect.height > 8 and lower_y <= rect.y1 <= limit_y:
            rects.append(fitz.Rect(rect))
    for image in page.get_images(full=True):
        xref = image[0]
        for rect in page.get_image_rects(xref):
            if rect.width > 40 and rect.height > 40 and lower_y <= rect.y1 <= limit_y:
                rects.append(fitz.Rect(rect))
    union = union_rects(rects)
    if union is None:
        lower_y = caption_rect.y1 + 4
        upper_y = min(page.rect.y1, caption_rect.y1 + 620)
        for drawing in page.get_drawings():
            rect = drawing.get("rect")
            if rect and rect.width > 8 and rect.height > 8 and lower_y <= rect.y0 <= upper_y:
                rects.append(fitz.Rect(rect))
        for image in page.get_images(full=True):
            xref = image[0]
            for rect in page.get_image_rects(xref):
                if rect.width > 40 and rect.height > 40 and lower_y <= rect.y0 <= upper_y:
                    rects.append(fitz.Rect(rect))
        union = union_rects(rects)
    if union is None:
        return None
    return expand_rect(union, page.rect, 18)


def build_rendered_image_blocks(pdf_path: Path, output_dir: Path) -> dict[int, list[ImageBlock]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    images_by_page: dict[int, list[ImageBlock]] = {}
    figure_index = 1
    with fitz.open(pdf_path) as doc:
        for page_index, page in enumerate(doc, start=1):
            captions = caption_rects(page)
            if captions:
                added_from_caption = False
                for figure_id, caption_rect in captions:
                    clip = figure_clip_for_caption(page, caption_rect)
                    if clip is None or clip.width < 40 or clip.height < 40:
                        continue
                    image_path = output_dir / f"page_{page_index:04d}_figure_{figure_id.replace('.', '_')}.png"
                    render_clip(page, clip, image_path)
                    images_by_page.setdefault(page_index, []).append(ImageBlock(image_path))
                    figure_index += 1
                    added_from_caption = True
                if added_from_caption:
                    continue
            candidates: list[fitz.Rect] = []
            for image in page.get_images(full=True):
                xref = image[0]
                for rect in page.get_image_rects(xref):
                    if rect.width < 120 or rect.height < 120:
                        continue
                    ratio = max(rect.width / max(rect.height, 1), rect.height / max(rect.width, 1))
                    if ratio > 8:
                        continue
                    candidates.append(fitz.Rect(rect))
            candidates.sort(key=lambda rect: (rect.y0, rect.x0))
            for image_index, rect in enumerate(candidates, start=1):
                image_path = output_dir / f"page_{page_index:04d}_image_{image_index:02d}.png"
                render_clip(page, expand_rect(rect, page.rect, 4), image_path)
                images_by_page.setdefault(page_index, []).append(ImageBlock(image_path))
    return images_by_page


def load_scanned_figure_blocks(manifest_path: Path) -> dict[int, list[ImageBlock]]:
    if not manifest_path.exists():
        return {}
    try:
        figures = json.loads(manifest_path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    images_by_page: dict[int, list[ImageBlock]] = {}
    for item in figures:
        try:
            page = int(item.get("page"))
            path = Path(str(item.get("path", "")))
        except Exception:
            continue
        if path.exists():
            images_by_page.setdefault(page, []).append(ImageBlock(path))
    return images_by_page


def sibling_scanned_figure_manifest(translated_dir: Path) -> Path:
    name = translated_dir.name
    if name.endswith("_translated"):
        figure_dir = translated_dir.parent / f"{name[:-len('_translated')]}_figures"
    else:
        figure_dir = translated_dir.parent / f"{name}_figures"
    return figure_dir / "figures_manifest.json"


def normalize_image_for_docx(path: Path) -> Path | None:
    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff"}:
        return path
    png_path = path.with_suffix(".png")
    try:
        with Image.open(path) as image:
            if image.mode not in {"RGB", "RGBA"}:
                image = image.convert("RGB")
            image.save(png_path, "PNG")
        return png_path
    except (UnidentifiedImageError, OSError):
        return None


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_width(draw: ImageDraw.ImageDraw, text: str, font) -> int:
    box = draw.textbbox((0, 0), text, font=font)
    return int(box[2] - box[0])


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill=(0, 0, 0)) -> None:
    x0, y0, x1, y1 = box
    lines = wrap_text(draw, text, font, max(10, x1 - x0 - 20))
    line_h = int(font.size * 1.25) if hasattr(font, "size") else 24
    total_h = line_h * len(lines)
    y = y0 + (y1 - y0 - total_h) / 2
    for line in lines:
        w = text_width(draw, line, font)
        draw.text((x0 + (x1 - x0 - w) / 2, y), line, font=font, fill=fill)
        y += line_h


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        current = ""
        for char in paragraph:
            candidate = current + char
            if current and text_width(draw, candidate, font) > max_width:
                lines.append(current)
                current = char
            else:
                current = candidate
        lines.append(current)
    return lines


def draw_wrapped_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font,
    max_width: int,
    fill=(0, 0, 0),
    line_spacing: float = 1.35,
) -> int:
    x, y = xy
    line_h = int((font.size if hasattr(font, "size") else 22) * line_spacing)
    for line in wrap_text(draw, text, font, max_width):
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h
    return y


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill=(194, 67, 43), width: int = 8) -> None:
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    dx, dy = ex - sx, ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 24
    p1 = (ex - ux * size + px * size * 0.55, ey - uy * size + py * size * 0.55)
    p2 = (ex - ux * size - px * size * 0.55, ey - uy * size - py * size * 0.55)
    draw.polygon([end, p1, p2], fill=fill)


def create_gilbert_figure_1_3(path: Path) -> None:
    width, height = 1500, 1200
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = pil_font(44, True)
    body_font = pil_font(34)
    small_font = pil_font(26)
    blue = (45, 82, 128)
    gray = (120, 120, 120)
    draw.text((80, 60), "教练员价值观工作表", font=title_font, fill=blue)
    prompts = [
        "作为教练员，我最喜欢自己的三点是什么？最不满意自己的三点是什么？",
        "我认识的最快乐的教练员是谁？",
        "我最喜欢、最尊重的两位教练员是谁？为什么？",
        "作为教练员，我是谁？",
    ]
    y = 160
    for prompt in prompts:
        y = draw_wrapped_text(draw, (80, y), prompt, body_font, width - 160)
        y += 26
        for _ in range(3):
            draw.line((80, y, width - 80, y), fill=(185, 185, 185), width=3)
            y += 58
        y += 22
    image.save(path)


def create_gilbert_figure_1_4(path: Path) -> None:
    width, height = 1500, 1850
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = pil_font(38, True)
    body_font = pil_font(30)
    header_font = pil_font(30, True)
    small_font = pil_font(24)
    gray = (120, 120, 120)
    line = (150, 150, 150)
    header_bg = (220, 225, 225)
    draw.text((70, 55), "执教使命（你为什么执教？）：作为一名教练员，我存在的意义是……", font=title_font, fill=(35, 35, 35))
    y = 130
    for _ in range(3):
        draw.line((70, y, width - 70, y), fill=line, width=3)
        y += 58
    y += 28
    left, right = 70, width - 70
    mid = 410
    row_h = 240
    header_h = 58
    draw.rectangle((left, y, right, y + header_h), fill=header_bg, outline=(90, 90, 90), width=3)
    draw.line((mid, y, mid, y + header_h + row_h * 5), fill=(120, 120, 120), width=3)
    draw_centered_text(draw, (left, y, mid, y + header_h), "核心价值观", header_font)
    draw_centered_text(draw, (mid, y, right, y + header_h), "核心价值观行动陈述", header_font)
    y += header_h
    for index in range(5):
        draw.rectangle((left, y, right, y + row_h), outline=(150, 150, 150), width=3)
        draw.line((mid, y, mid, y + row_h), fill=(150, 150, 150), width=3)
        draw.text((left + 24, y + 22), str(index + 1), font=body_font, fill=(55, 55, 55))
        y += row_h
    image.save(path)


def create_gilbert_figure_3_3(path: Path) -> None:
    width, height = 1600, 680
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = pil_font(32)
    bold = pil_font(34, True)
    small = pil_font(28)
    bg = (216, 238, 239)
    box_bg = (226, 232, 244)
    box_outline = (142, 154, 176)
    red = (197, 70, 45)
    draw.rounded_rectangle((70, 65, width - 70, height - 65), radius=16, fill=bg, outline=(75, 110, 112), width=4)
    draw_centered_text(draw, (70, 72, width - 70, 118), "运动情境", bold)
    draw_centered_text(draw, (70, height - 118, width - 70, height - 72), "运动情境", bold)

    left_boxes = [
        (110, 145, 410, 235, "运动员目标"),
        (110, 260, 410, 350, "运动员情绪"),
        (110, 375, 410, 485, "运动员自我能力感知"),
    ]
    right_boxes = [
        (1190, 145, 1490, 235, "教练员目标"),
        (1190, 260, 1490, 350, "教练员情绪"),
        (1190, 375, 1490, 485, "教练员对运动员能力的感知"),
    ]
    for x0, y0, x1, y1, label in left_boxes + right_boxes:
        draw.rectangle((x0, y0, x1, y1), fill=box_bg, outline=box_outline, width=4)
        draw_centered_text(draw, (x0, y0, x1, y1), label, font)

    draw.ellipse((500, 140, 1100, 500), fill=(250, 220, 205), outline=(180, 160, 150), width=3)
    draw_centered_text(draw, (650, 255, 950, 360), "通过互动\n形成理解", bold)
    draw.text((555, 205), "提问", font=small, fill=(55, 55, 55))
    draw.text((930, 205), "倾听", font=small, fill=(55, 55, 55))
    draw.text((925, 405), "观察", font=small, fill=(55, 55, 55))
    draw.text((565, 405), "时间", font=small, fill=(55, 55, 55))
    draw_arrow(draw, (660, 190), (915, 190), fill=red)
    draw_arrow(draw, (1015, 250), (1015, 390), fill=red)
    draw_arrow(draw, (925, 450), (665, 450), fill=red)
    draw_arrow(draw, (585, 390), (585, 250), fill=red)
    image.save(path)


def apply_gilbert_figure_overrides(images_by_page: dict[int, list[ImageBlock]], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    figure_specs = {
        34: ("gilbert_fig_1_3.png", create_gilbert_figure_1_3),
        36: ("gilbert_fig_1_4.png", create_gilbert_figure_1_4),
        96: ("gilbert_fig_3_3.png", create_gilbert_figure_3_3),
    }
    for page, (filename, creator) in figure_specs.items():
        path = output_dir / filename
        creator(path)
        images_by_page[page] = [ImageBlock(path)]
    # The original English worksheet artwork for figures 1.3 and 1.4 lives on the following pages.
    # The rebuilt Chinese figures above replace those images, so the English originals are suppressed.
    for page in [35, 37]:
        images_by_page[page] = []


def is_reference_start(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.lower().startswith(("in:", "and ", "of ", "the ", "than ", "with ")):
        return False
    if re.match(r"^\d{1,3}[\t .]+[A-ZÀ-ÖØ-Þ]", stripped):
        return True
    return bool(re.search(r"\(\s*\d{4}[a-z]?\s*\)", stripped[:140]))


def clean_reference_source_line(line: str) -> str:
    line = line.replace("\u00ad", "")
    line = re.sub(r"\s+", " ", line)
    return line.strip()


def is_reference_noise_line(line: str) -> bool:
    upper = line.upper()
    if "COPYRIGHT" in upper and ("\u0411\u0418\u0411\u041a\u041e\u041c" in upper or "K\u041d\u0418\u0413\u0410-C\u0415\u0420\u0412\u0418\u0421" in upper):
        return True
    if "\u0414\u0412\u0418\u0413\u0410\u0422\u0415\u041b\u042c\u041d\u042b\u0415" in upper and "\u0421\u041f\u041e\u0420\u0422\u0421\u041c\u0415\u041d\u041e\u0412" in upper:
        return True
    return is_running_header_text(line)


def source_reference_text_to_blocks(text: str, include_heading: bool) -> list[Block]:
    lines = [clean_reference_source_line(line) for line in text.splitlines() if clean_reference_source_line(line)]
    lines = [line for line in lines if not re.fullmatch(r"\d+|[ivxlcdm]+", line, flags=re.I)]
    lines = [line for line in lines if not is_reference_noise_line(line)]
    blocks: list[Block] = []
    current: list[str] = []
    in_reference_list = not include_heading

    def flush_current() -> None:
        nonlocal current
        if current:
            blocks.append(ParagraphBlock(" ".join(current), "reference"))
            current = []

    def append_reference_line(line: str) -> None:
        if current and re.search(r"[-–]$", current[-1]):
            current[-1] = re.sub(r"[-–]$", "", current[-1]) + line
        else:
            current.append(line)

    for line in lines:
        heading = REFERENCE_HEADING_RE.search(line)
        bibliography_heading = re.fullmatch(r"Bibliography", line, flags=re.I)
        works_cited_heading = re.fullmatch(r"Works\s+Cited", line, flags=re.I)
        citation_list_heading = re.fullmatch(r"List\s+of\s+Citations.*", line, flags=re.I)
        russian_heading = RUSSIAN_REFERENCE_HEADING_RE.fullmatch(line) is not None
        if heading or re.fullmatch(r"References", line, flags=re.I) or bibliography_heading or works_cited_heading or citation_list_heading or russian_heading:
            if in_reference_list:
                continue
            current = []
            if include_heading:
                if bibliography_heading:
                    heading_text = "Bibliography"
                elif works_cited_heading:
                    heading_text = "Works Cited"
                elif citation_list_heading:
                    heading_text = line
                elif russian_heading:
                    heading_text = "参考文献"
                else:
                    heading_text = f"References for Chapter {heading.group(1)}" if heading else "References"
                blocks.append(ParagraphBlock(heading_text, "h1" if bibliography_heading else "h2"))
            in_reference_list = True
            continue
        if not in_reference_list:
            continue
        if re.fullmatch(r"(?:Introduction|Chapter\s+\d+|Conclusion)", line, flags=re.I):
            flush_current()
            blocks.append(ParagraphBlock(line, "h3"))
            continue
        if is_reference_start(line):
            if current:
                flush_current()
            current = [line]
        elif current:
            append_reference_line(line)
        else:
            current = [line]
    if current:
        flush_current()
    return blocks


def source_reference_blocks_by_page(pdf_path: Path) -> dict[int, list[Block]]:
    refs: dict[int, list[Block]] = {}
    in_refs = False
    known_content_starts = (
        set(SOCCER_CHAPTER_START_PAGES)
        | set(GOMES_CHAPTER_START_PAGES)
        | set(GILBERT_CHAPTER_START_PAGES)
        | set(GORDON_CHAPTER_START_PAGES)
        | set(LTAD_CHAPTER_START_PAGES)
        | set(STRENGTH_CHAPTER_START_PAGES)
        | set(PLATONOV_MOTOR_CHAPTER_START_PAGES)
        | set(PLATONOV_MOTOR_PART_START_PAGES)
        | set(BONDARCHUK_CHAPTER_START_PAGES)
        | set(SCIENCE_SOCCER_CHAPTER_START_PAGES)
        | set(SCIENCE_SOCCER_SECTION_TITLES)
        | set(SOVIET_WEIGHTLIFTING_START_PAGES)
    )
    with fitz.open(pdf_path) as doc:
        for page_index, page in enumerate(doc, start=1):
            text = repair_pdf_glyph_text(page.get_text("text"))
            all_lines = [line.strip() for line in text.splitlines() if line.strip()]
            top_lines = all_lines[:5]
            starts_about_author = bool(re.search(r"(^|\n)\s*About\s+the\s+Author\s*(\n|$)", text, flags=re.I))
            starts_glossary = bool(re.search(r"(^|\n)\s*Glossary\s*(\n|$)", text, flags=re.I))
            starts_index = bool(re.search(r"(^|\n)\s*Index\s*(\n|$)", text, flags=re.I))
            starts_chapter = any(re.fullmatch(r"CHAPTER\s+\d+", line) for line in top_lines)
            starts_new_section = starts_about_author or starts_glossary or starts_index or starts_chapter
            has_ref_heading = any(
                REFERENCE_HEADING_RE.search(line) is not None
                or re.fullmatch(r"References|Bibliography|Works\s+Cited|List\s+of\s+Citations.*", line, flags=re.I)
                or RUSSIAN_REFERENCE_HEADING_RE.fullmatch(line) is not None
                for line in all_lines
            )
            looks_like_front_toc = page_index <= 20 and (
                len(re.findall(r"\bChapter\s+\d+", text, flags=re.I)) >= 2
                or len(re.findall(r"\bPart\s+[IVX]+", text, flags=re.I)) >= 2
                or re.search(r"(^|\n)\s*Contents\s*(\n|$)", text, flags=re.I) is not None
            )
            if looks_like_front_toc:
                has_ref_heading = False
            if page_index in known_content_starts and not has_ref_heading:
                in_refs = False
            if starts_new_section and not has_ref_heading:
                in_refs = False
            was_in_refs = in_refs
            if has_ref_heading:
                in_refs = True
            if in_refs:
                refs[page_index] = source_reference_text_to_blocks(text, include_heading=bool(has_ref_heading and not was_in_refs))
    return refs


def run_final_pipeline(input_pdf: Path, translated_dir: Path, output_dir: Path, title: str) -> dict:
    clean_output_dir(output_dir)
    with tempfile.TemporaryDirectory(prefix="book_final_") as tmp:
        tmp_path = Path(tmp)
        source_pages = source_text_by_page(input_pdf)
        scanned_figure_manifest = sibling_scanned_figure_manifest(translated_dir)
        if scanned_figure_manifest.exists():
            images_by_page = load_scanned_figure_blocks(scanned_figure_manifest)
        elif is_image_only_pdf(source_pages):
            images_by_page = {}
        else:
            images_by_page = build_rendered_image_blocks(input_pdf, tmp_path / "images")
        reference_blocks = {} if is_image_only_pdf(source_pages) else source_reference_blocks_by_page(input_pdf)
        use_standard_sections = "高级运动训练" in title or "issurin" in input_pdf.name.lower()
        lower_name = input_pdf.name.lower()
        use_weightlifting_chapter_titles = "中国举重" in title or "chinese_weightlifting" in lower_name
        use_soccer_chapter_titles = "足球" in title or "soccer" in lower_name
        use_frank_chapter_titles = "运动训练原则" in title or "sports_training_principles" in lower_name
        use_gomes_chapter_titles = "结构化与周期化" in title or "gomes_treinamento" in lower_name
        use_gilbert_chapter_titles = "执教" in title or "gilbert_coaching" in lower_name or "coaching_better" in lower_name
        use_gordon_chapter_titles = "教练科学" in title or "gordon" in lower_name or "coaching_science" in lower_name
        use_ltad_chapter_titles = "长期运动员发展" in title or "ltad" in lower_name or "long-term_athlete_development" in lower_name
        use_strength_chapter_titles = (
            "力量训练" in title
            or "science_practice_strength_training" in lower_name
            or "science_and_practice_of_strength_training" in lower_name
        )
        use_platonov_chapter_titles = "运动训练分期理论" in title or "periodization_theory" in lower_name or "periodiz" in lower_name
        use_platonov_motor_chapter_titles = (
            "运动员的运动能力" in title
            or "运动员的运动素质" in title
            or "platonov_motor" in lower_name
            or "motor_abilities" in lower_name
        )
        use_bondarchuk_chapter_titles = (
            "高水平运动员训练过程管理" in title
            or "训练过程管理" in title
            or "bondarchuk" in lower_name
            or "manage_training_process" in lower_name
        )
        use_science_soccer_chapter_titles = (
            "足球科学" in title
            or "science_and_soccer" in lower_name
            or "science and soccer" in lower_name
        )
        use_soviet_weightlifting_titles = (
            "苏联举重体系" in title
            or "soviet_sports_training" in lower_name
            or "soviet_weightlifting" in lower_name
            or "pdfcoffee.com_soviet" in lower_name
        )
        chapter_title_overrides = None
        if use_soviet_weightlifting_titles:
            chapter_title_overrides = SOVIET_WEIGHTLIFTING_TITLES
        elif use_science_soccer_chapter_titles:
            chapter_title_overrides = SCIENCE_SOCCER_CHAPTER_TITLES
        elif use_bondarchuk_chapter_titles:
            chapter_title_overrides = BONDARCHUK_CHAPTER_TITLES
        elif use_platonov_motor_chapter_titles:
            chapter_title_overrides = PLATONOV_MOTOR_CHAPTER_TITLES
        elif use_platonov_chapter_titles:
            chapter_title_overrides = PLATONOV_CHAPTER_TITLES
        elif use_weightlifting_chapter_titles:
            chapter_title_overrides = CHINESE_WEIGHTLIFTING_CHAPTER_TITLES
        elif use_soccer_chapter_titles:
            chapter_title_overrides = SOCCER_CHAPTER_TITLES
        elif use_frank_chapter_titles:
            chapter_title_overrides = FRANK_CHAPTER_TITLES
        elif use_gomes_chapter_titles:
            chapter_title_overrides = GOMES_CHAPTER_TITLES
        elif use_gilbert_chapter_titles:
            chapter_title_overrides = GILBERT_CHAPTER_TITLES
            apply_gilbert_figure_overrides(images_by_page, tmp_path / "gilbert_rebuilt_figures")
        elif use_gordon_chapter_titles:
            chapter_title_overrides = GORDON_CHAPTER_TITLES
        elif use_ltad_chapter_titles:
            chapter_title_overrides = LTAD_CHAPTER_TITLES
        elif use_strength_chapter_titles:
            chapter_title_overrides = STRENGTH_CHAPTER_TITLES
        blocks = parse_translated_markdown(
            translated_dir,
            images_by_page,
            reference_blocks,
            source_pages,
            use_standard_sections=use_standard_sections,
            chapter_title_overrides=chapter_title_overrides,
        )
        docx_path = output_dir / "final_zh.docx"
        pdf_path = output_dir / "final_zh.pdf"
        render_docx(blocks, docx_path, title)
        export_pdf_with_word(docx_path, pdf_path)
    errors = validate_final_output(output_dir)
    if errors:
        raise RuntimeError("; ".join(errors))
    return {
        "output_dir": str(output_dir),
        "docx": str(output_dir / "final_zh.docx"),
        "pdf": str(output_dir / "final_zh.pdf"),
        "blocks": len(blocks),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate final Chinese DOCX and PDF only.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--translated-dir", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--title", required=True)
    args = parser.parse_args()
    result = run_final_pipeline(args.input.resolve(), args.translated_dir.resolve(), args.output_dir.resolve(), args.title)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
